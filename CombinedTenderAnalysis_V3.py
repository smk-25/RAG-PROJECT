# CombinedTenderAnalysis_V3.py
# Modernized Version with Professional UI/UX
import os
import re
import gc
import time
import json
import uuid
import hashlib
import importlib
import pickle
import tempfile
import asyncio
import random
import collections
import base64
from io import BytesIO
from gtts import gTTS
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import streamlit as st
import numpy as np
import pandas as pd
import chromadb
import fitz  # PyMuPDF
import nltk
from google import genai
from docx import Document
import pdfplumber
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# --------------------------
# Page Configuration
# --------------------------
st.set_page_config(
    page_title="Tender Analyzer Pro V2",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --------------------------
# Custom CSS Design System
# --------------------------
st.markdown("""
    <style>
    /* Main Styling */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    
    .main {
        background-color: #f8f9fa;
    }
    
    /* Card Styling */
    .stCard {
        background-color: white;
        padding: 1.5rem;
        border-radius: 0.75rem;
        border: 1px solid #e9ecef;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        margin-bottom: 1.5rem;
    }
    
    .metric-card {
        text-align: center;
        padding: 1rem;
        background: #ffffff;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        border: 1px solid #eee;
    }
    
    /* Header Styling */
    .app-header {
        background: linear-gradient(90deg, #1e3a8a 0%, #3b82f6 100%);
        padding: 2rem;
        border-radius: 0.75rem;
        color: white;
        margin-bottom: 2rem;
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
    }
    
    /* Success/Info/Warning overrides */
    div[data-testid="stMetricValue"] {
        font-size: 1.8rem;
        font-weight: 700;
        color: #1e3a8a;
    }
    
    /* Dial Gauge Styling */
    .gauge-container {
        width: 250px;
        margin: 20px auto;
        position: relative;
        text-align: center;
    }
    .gauge-wrap {
        width: 200px;
        height: 100px;
        position: relative;
        margin: 0 auto;
        overflow: hidden;
    }
    .gauge-bg {
        width: 200px;
        height: 200px;
        border-radius: 50%;
        background: #e5e7eb;
        position: absolute;
    }
    .gauge-fill {
        width: 200px;
        height: 200px;
        border-radius: 50%;
        background: conic-gradient(from -90deg at 50% 50%, #10B981 0deg, #10B981 var(--angle), transparent var(--angle));
        position: absolute;
        transition: transform 1s ease-in-out;
    }
    .gauge-inner {
        width: 160px;
        height: 80px;
        background: #ffffff;
        border-top-left-radius: 80px;
        border-top-right-radius: 80px;
        position: absolute;
        bottom: 0;
        left: 20px;
        display: flex;
        align-items: flex-end;
        justify-content: center;
        padding-bottom: 5px;
        font-weight: 700;
        font-size: 1.5rem;
    }
    .legend-container {
        display: flex;
        justify-content: center;
        gap: 15px;
        margin-top: 10px;
        font-size: 0.75rem;
    }
    .legend-item { display: flex; align-items: center; gap: 5px; }
    .legend-color { width: 12px; height: 12px; border-radius: 2px; }

    /* Sidebar Styling */
    .css-1d391kg {
        background-color: #f1f5f9;
    }
    </style>
    """, unsafe_allow_html=True)

# --------------------------
# Optional & Heavy Imports
# --------------------------
try:
    from sentence_transformers import CrossEncoder
    CROSS_ENCODER_AVAILABLE = True
except Exception:
    CrossEncoder = None
    CROSS_ENCODER_AVAILABLE = False
try:
    from langchain_groq import ChatGroq
    from langchain.schema import HumanMessage
    LANGCHAIN_AVAILABLE = True
except Exception:
    ChatGroq = HumanMessage = None
    LANGCHAIN_AVAILABLE = False
try:
    from bert_score import score as bert_score_fn
    BERTSCORE_AVAILABLE = True
except Exception:
    bert_score_fn = None
    BERTSCORE_AVAILABLE = False
try:
    from rank_bm25 import BM25Okapi
    BM25_AVAILABLE = True
except Exception:
    BM25Okapi = None
    BM25_AVAILABLE = False
try:
    from langchain_community.document_loaders import PyPDFLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_PDF_AVAILABLE = True
except Exception:
    PyPDFLoader = None
    RecursiveCharacterTextSplitter = None
    LANGCHAIN_PDF_AVAILABLE = False
try:
    import tiktoken as _tiktoken
except Exception:
    _tiktoken = None

# --------------------------
# Constants
# --------------------------
MAX_FILE_SIZE_MB = 50
MAX_QUERY_LENGTH = 5000
DEFAULT_RATE_LIMIT_SLEEP = 60.0
ZOOM_LEVEL = 1.5
MAX_CHUNK_TOKENS = 4096
DEFAULT_EMBEDDING_DIM = 384
CHARS_PER_TOKEN_ESTIMATE = 4
RRF_RANK_CONSTANT = 60
HYBRID_FUSION_RRF_WEIGHT = 0.5
HYBRID_FUSION_DENSE_WEIGHT = 0.3
HYBRID_FUSION_SPARSE_WEIGHT = 0.2
CROSS_ENCODER_WEIGHT = 0.6
ORIGINAL_SCORE_WEIGHT = 0.4
MULTI_QUERY_RRF_WEIGHT = 0.6
MULTI_QUERY_AVG_WEIGHT = 0.4
MIN_CLAUSE_LENGTH = 50
ANSWER_VALIDATION_CONTEXT_CHARS = 2000
MAX_CLAUSE_DISPLAY_LENGTH = 200
MAX_CHUNK_PREVIEW_LENGTH = 300
DENSE_RETRIEVAL_CANDIDATES = 100
DEFAULT_MISSING_RANK = 1000
MAX_QUERY_VARIATIONS = 2
MULTI_QUERY_RETRIEVAL_MULTIPLIER = 2
CROSS_ENCODER_CANDIDATE_MULTIPLIER = 3
CLAUSE_PATTERNS = [
    r'\d+\.\d+(?:\.\d+)?',
    r'[A-Z][A-Za-z\s]+:',
    r'\([a-z]\)',
    r'[•\-]\s+',
]
EMBEDDING_BATCH_SIZE = 64          # Encode in batches to limit peak memory
MAX_CITATION_SENTENCES = 50        # Cap sentences fed into citation embedding pass
MAX_CHUNKS_WARNING_THRESHOLD = 1500  # Warn user when chunk count is very high
REDUCE_BATCH_THRESHOLD = 50        # Max findings before switching to hierarchical reduce
# Common English stop-words excluded from evidence chunk-search scoring so that
# low-value tokens (e.g. "the", "and") do not dominate the relevance score.
_EVIDENCE_STOP_WORDS = frozenset({
    "the", "a", "an", "of", "in", "to", "for", "and", "or", "is",
    "are", "will", "shall", "be", "at", "by",
})

# --------------------------
# Initialization
# --------------------------
@st.cache_resource
def setup_nltk_shared():
    try:
        nltk.download('punkt', quiet=True)
        nltk.download('punkt_tab', quiet=True)
        return True
    except Exception as e:
        st.warning(f"Failed to download NLTK data: {e}")
        return False
_NLTK_READY = setup_nltk_shared()

@st.cache_resource
def load_embedding_model_cached(model_name: str = "all-MiniLM-L6-v2"):
    """Load (and cache) a SentenceTransformer so the model is only downloaded
    and allocated once per Streamlit server lifetime, regardless of how many
    times the user presses 'Index Documents'.
    """
    return SentenceTransformer(model_name)

@st.cache_resource
def load_cross_encoder_cached(model_name: str):
    """Load (and cache) a CrossEncoder model. Returns None when the
    sentence-transformers CrossEncoder is unavailable.
    """
    if not CROSS_ENCODER_AVAILABLE or CrossEncoder is None:
        return None
    try:
        return CrossEncoder(model_name)
    except Exception:
        return None

@st.cache_resource
def _get_em(name: str):
    """Return a cached EmbeddingManagerRAG for the given model name.

    Defined at module level so @st.cache_resource can correctly identify and
    reuse the cached resource across Streamlit reruns.  Defining this inside a
    conditional block caused a new function object (and therefore a new cache
    entry and a new ChromaDB client) to be created on every rerun, leading to
    SQLite database-lock errors during document indexing.
    """
    return EmbeddingManagerRAG(name)

@st.cache_resource
def _get_vs():
    """Return a cached VectorStoreRAG instance.

    Must live at module level for the same reason as _get_em above.
    """
    return VectorStoreRAG()

# Shared Utils
def sha256_text_shared(text: str) -> str:
    return hashlib.sha256((text or "").strip().encode("utf-8")).hexdigest()

def sha256_file_shared(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""): h.update(chunk)
    return h.hexdigest()

def normalize_text_shared(s: str) -> str:
    if s is None: return ""
    s = s.lower()
    s = re.sub(r"[^a-z0-9\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def tokenize_simple_shared(s: str) -> List[str]:
    return re.findall(r"\w+", (s or "").lower())

_SENTENCE_RE = re.compile(r'(?<=[.!?])\s+')
def simple_sent_tokenize_shared(text: str) -> List[str]:
    sents = [s.strip() for s in _SENTENCE_RE.split(text or "") if s.strip()]
    return sents if sents else ([text.strip()] if (text or "").strip() else [])

_PROV_SPLIT_RE = re.compile(r"\n\s*(\[\d+\]\s*Source:|→\s*Source:|Source:|Source\s*[:\-])", flags=re.IGNORECASE)
def strip_provenance(text: str) -> str:
    if not text: return ""
    parts = _PROV_SPLIT_RE.split(text)
    cleaned = parts[0] if parts else text
    cleaned = re.sub(r"\(chunk:?[^\)]*\)", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"[\[\(]\s*\d+\s*[\]\)]", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned

def compute_mean_support_score_shared(results: List[Dict[str, Any]]) -> float:
    if not results: return 0.0
    scores = [float(r.get("similarity_score", 0.0)) for r in results]
    return float(np.max(scores)) if scores else 0.0

# --- Evaluation Metrics ---
def token_precision_shared(pred: str, gold: str) -> float:
    p_tokens = tokenize_simple_shared(pred)
    g_tokens = tokenize_simple_shared(gold)
    if not p_tokens: return 0.0
    common = collections.Counter(p_tokens) & collections.Counter(g_tokens)
    return sum(common.values()) / len(p_tokens)

def token_recall_shared(pred: str, gold: str) -> float:
    p_tokens = tokenize_simple_shared(pred)
    g_tokens = tokenize_simple_shared(gold)
    if not g_tokens: return 1.0
    common = collections.Counter(p_tokens) & collections.Counter(g_tokens)
    return sum(common.values()) / len(g_tokens)

def rouge_l_shared(pred: str, gold: str) -> float:
    p, g = tokenize_simple_shared(pred), tokenize_simple_shared(gold)
    if not p or not g: return 0.0
    n, m = len(p), len(g)
    prev = [0] * (m + 1)
    for i in range(1, n + 1):
        curr = [0] * (m + 1)
        for j in range(1, m + 1):
            if p[i-1] == g[j-1]: curr[j] = prev[j-1] + 1
            else: curr[j] = max(prev[j], curr[j-1])
        prev = curr
    lcs = prev[m]
    prec, rec = lcs/n, lcs/m
    return 2*prec*rec/(prec+rec) if (prec+rec) > 0 else 0.0

# --------------------------
# PDF Loading & Chunking
# --------------------------
class RAGDocument:
    def __init__(self, page_content: str, metadata: dict = None):
        self.page_content = page_content
        self.metadata = metadata or {}

def load_pdf_with_pymupdf(pdf_path: str) -> List[RAGDocument]:
    """
    Extract full text from every PDF page using:
    - PyMuPDF block-sorted extraction (preserves reading order for multi-column layouts)
    - pdfplumber table detection (appends tables as Markdown so no content is lost)
    - Text cleaning (null bytes, redundant whitespace)
    Mirrors the extract_pages() approach from Summarizationcode.py.
    """
    docs = []
    try:
        pdf_document = fitz.open(pdf_path)
        with pdfplumber.open(pdf_path) as pl_pdf:
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]

                # Block-sorted extraction preserves correct reading order for
                # multi-column and complex layouts.
                try:
                    blocks = page.get_text("blocks")
                    blocks = sorted(blocks, key=lambda b: (round(b[1], 1), round(b[0], 1)))
                    text = "\n".join(
                        b[4] for b in blocks if isinstance(b, (list, tuple)) and len(b) >= 5
                    )
                except Exception:
                    text = page.get_text()

                # Extract tables via pdfplumber and append as cleaned Markdown so
                # tabular evidence is fully captured and not silently dropped.
                if page_num < len(pl_pdf.pages):
                    tables = pl_pdf.pages[page_num].extract_tables()
                    for table in tables:
                        if table:
                            df_t = pd.DataFrame(table).dropna(how="all").dropna(axis=1, how="all")
                            if not df_t.empty:
                                text += (
                                    f"\n\n[TABLE DATA - PAGE {page_num + 1}]:\n"
                                    + df_t.to_markdown(index=False)
                                    + "\n\n"
                                )

                # Normalise whitespace and remove null bytes that can corrupt
                # downstream embeddings.
                text = text.replace("\x00", " ").replace("\r\n", "\n")
                text = re.sub(r"[ \t]+", " ", text)
                text = re.sub(r"\n{3,}", "\n\n", text).strip()

                if text:
                    docs.append(RAGDocument(page_content=text, metadata={"page": page_num + 1}))
        pdf_document.close()
    except Exception as e:
        st.error(f"Error loading PDF: {e}")
    return docs

def process_pdf(pdf_path: str) -> List[Any]:
    """Load a PDF using PyPDFLoader (mirrors EnhancedRAG10.3.py logic).

    Adds source_file, file_type, and file_hash metadata to every page document.
    Falls back to load_pdf_with_pymupdf when PyPDFLoader / langchain is not
    available in the environment.
    """
    if not LANGCHAIN_PDF_AVAILABLE or PyPDFLoader is None:
        st.write(f"Loading PDF: {Path(pdf_path).name}")
        return load_pdf_with_pymupdf(pdf_path)
    st.write(f"Loading PDF: {Path(pdf_path).name}")
    try:
        loader = PyPDFLoader(pdf_path)
        docs = loader.load()
    except Exception as e:
        st.error(f"PDF loader error: {e}")
        return []
    try:
        file_hash = sha256_file_shared(pdf_path)
    except Exception:
        file_hash = None
    for d in docs:
        d.metadata["source_file"] = Path(pdf_path).name
        d.metadata["file_type"] = "pdf"
        if file_hash:
            d.metadata["file_hash"] = file_hash
    st.write(f"Loaded {len(docs)} pages")
    return docs

def chunk_recursive_character(docs: List[RAGDocument], chunk_size: int = 800, chunk_overlap: int = 128) -> List[RAGDocument]:
    chunks = []
    for doc in docs:
        text = doc.page_content
        metadata = doc.metadata.copy()
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            if chunk_text.strip():
                chunks.append(RAGDocument(page_content=chunk_text, metadata=metadata))
            start += chunk_size - chunk_overlap
            if start >= len(text): break
    return chunks

def chunk_hybrid_token_semantic(docs: List[RAGDocument], chunk_size: int = 800, chunk_overlap: int = 128) -> List[RAGDocument]:
    chunks = []
    for doc in docs:
        text = doc.page_content
        metadata = doc.metadata.copy()
        try:
            sentences = nltk.sent_tokenize(text) if _NLTK_READY else simple_sent_tokenize_shared(text)
        except: sentences = simple_sent_tokenize_shared(text)
        current_chunk, current_length = [], 0
        for sentence in sentences:
            sentence_length = len(sentence)
            if current_length + sentence_length > chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                if chunk_text.strip(): chunks.append(RAGDocument(page_content=chunk_text, metadata=metadata))
                overlap_sentences, overlap_length = [], 0
                for s in reversed(current_chunk):
                    if overlap_length + len(s) <= chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_length += len(s)
                    else: break
                current_chunk, current_length = overlap_sentences + [sentence], overlap_length + sentence_length
            else:
                current_chunk.append(sentence)
                current_length += sentence_length
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            if chunk_text.strip(): chunks.append(RAGDocument(page_content=chunk_text, metadata=metadata))
    return chunks

def chunk_fixed_context_window(docs: List[RAGDocument], window_size: int = 800) -> List[RAGDocument]:
    chunks = []
    for doc in docs:
        text = doc.page_content
        metadata = doc.metadata.copy()
        for i in range(0, len(text), window_size):
            chunk_text = text[i:i + window_size]
            if chunk_text.strip(): chunks.append(RAGDocument(page_content=chunk_text, metadata=metadata))
    return chunks

def apply_chunking_method(docs: List[RAGDocument], method: str, chunk_size: int = 800, chunk_overlap: int = 128) -> List[RAGDocument]:
    if method == "Recursive Character Splitter": return chunk_recursive_character(docs, chunk_size, chunk_overlap)
    elif method == "Hybrid (Token+Semantic)": return chunk_hybrid_token_semantic(docs, chunk_size, chunk_overlap)
    elif method == "Fixed Context Window": return chunk_fixed_context_window(docs, chunk_size)
    return chunk_recursive_character(docs, chunk_size, chunk_overlap)

def split_documents(documents: List[Any], chunk_size: int = 800, chunk_overlap: int = 128, method: str = "Recursive") -> List[Any]:
    """Split documents using RecursiveCharacterTextSplitter (mirrors EnhancedRAG10.3.py logic).

    Supports method names "Recursive", "Hybrid", and "Fixed-size" (matching
    EnhancedRAG10.3.py).  Falls back to apply_chunking_method when the langchain
    text splitter is unavailable.
    """
    if not LANGCHAIN_PDF_AVAILABLE or RecursiveCharacterTextSplitter is None:
        _method_map = {
            "Recursive": "Recursive Character Splitter",
            "Hybrid": "Hybrid (Token+Semantic)",
            "Fixed-size": "Fixed Context Window",
        }
        return apply_chunking_method(
            documents,  # type: ignore[arg-type]
            _method_map.get(method, "Recursive Character Splitter"),
            chunk_size,
            chunk_overlap,
        )

    # Token-aware length function: uses tiktoken when available, else char/4 estimate
    if _tiktoken is not None:
        try:
            _enc = _tiktoken.get_encoding("cl100k_base")
            def length_fn(text: str) -> int:
                return len(_enc.encode(text))
        except Exception:
            def length_fn(text: str) -> int:
                return max(1, len(text) // 4)
    else:
        def length_fn(text: str) -> int:
            return max(1, len(text) // 4)

    if method == "Recursive":
        st.write(f"Using Recursive chunking strategy (chunk_size={chunk_size}, overlap={chunk_overlap})")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=length_fn,
            separators=["\n\n", "\n", " ", ""]
        )
        split_docs = splitter.split_documents(documents)
        st.write(f"Split into {len(split_docs)} chunks using Recursive method")
        return split_docs

    elif method == "Hybrid":
        # Two-pass: first split on paragraph/line boundaries (semantic), then
        # further split any oversized chunks with overlap (fixed-size).
        st.write(f"Using Hybrid chunking strategy (chunk_size={chunk_size}, overlap={chunk_overlap})")
        semantic_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size * 2,
            chunk_overlap=0,
            length_function=length_fn,
            separators=["\n\n", "\n"]
        )
        semantic_chunks = semantic_splitter.split_documents(documents)
        final_chunks: List[Any] = []
        for doc in semantic_chunks:
            if length_fn(doc.page_content) > chunk_size:
                sub_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    length_function=length_fn,
                    separators=[" ", ""]
                )
                final_chunks.extend(sub_splitter.split_documents([doc]))
            else:
                final_chunks.append(doc)
        st.write(f"Split into {len(final_chunks)} chunks using Hybrid method")
        return final_chunks

    elif method == "Fixed-size":
        st.write(f"Using Fixed-size chunking strategy (chunk_size={chunk_size}, overlap={chunk_overlap})")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=length_fn,
            separators=[""]
        )
        split_docs = splitter.split_documents(documents)
        st.write(f"Split into {len(split_docs)} chunks using Fixed-size method")
        return split_docs

    else:
        st.warning(f"Unknown chunking method '{method}', defaulting to Recursive")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=length_fn,
            separators=["\n\n", "\n", " ", ""]
        )
        split_docs = splitter.split_documents(documents)
        st.write(f"Split into {len(split_docs)} chunks")
        return split_docs

# --------------------------
# RAG Core Components
# --------------------------
class EmbeddingManagerRAG:
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model_name = model_name
        # Reuse a cached model instance to avoid reloading weights on every run
        self.model = load_embedding_model_cached(model_name)
    def generate_embeddings(self, texts: List[str]) -> np.ndarray:
        if not texts:
            return np.zeros((0, DEFAULT_EMBEDDING_DIM), dtype=np.float32)
        # Encode in fixed-size batches to keep peak memory usage bounded
        return self.model.encode(texts, show_progress_bar=False, batch_size=EMBEDDING_BATCH_SIZE)

class VectorStoreRAG:
    def __init__(self, collection_name: str = "pdf_documents", path: str = "./vector_store"):
        self.client = chromadb.PersistentClient(path=path)
        self.collection = self.client.get_or_create_collection(name=collection_name)
    def add_documents(self, docs, embs):
        ids = [f"{d.metadata.get('file_hash','na')}_{sha256_text_shared(d.page_content)}" for d in docs]
        try:
            self.collection.delete(ids=ids)
        except Exception:
            pass
        # Deduplicate within the batch to avoid DuplicateIDError when chunks share identical content
        id_to_doc_emb: dict = {}
        for id_, doc, emb in zip(ids, docs, embs.tolist()):
            id_to_doc_emb[id_] = (doc, emb)
        unique_ids = list(id_to_doc_emb.keys())
        unique_docs, unique_embs = zip(*[id_to_doc_emb[i] for i in unique_ids]) if unique_ids else ([], [])
        self.collection.add(ids=unique_ids, documents=[d.page_content for d in unique_docs], metadatas=[d.metadata for d in unique_docs], embeddings=list(unique_embs))

class SparseBM25IndexRAG:
    def __init__(self): self.bm25, self.docs = None, []
    def build(self, docs):
        # Store only id/content/metadata – embeddings are already persisted in
        # ChromaDB and holding a second copy here wastes significant memory.
        self.docs = [{"id": sha256_text_shared(d.page_content), "content": d.page_content, "metadata": d.metadata} for d in docs]
        self.bm25 = BM25Okapi([tokenize_simple_shared(d.page_content) for d in docs])
    def is_ready(self): return self.bm25 is not None
    def get_top_n(self, q, n):
        if not self.bm25: return []
        scores = self.bm25.get_scores(tokenize_simple_shared(q))
        idx = np.argsort(scores)[::-1][:n]
        return [dict(self.docs[i], sparse_score=float(scores[i])) for i in idx]

class RAGRetrieverRAG:
    def __init__(self, vs, em, cross=None, sparse=None, use_s=False, s_k=200):
        self.vs, self.em, self.cross, self.sparse, self.use_s, self.s_k = vs, em, cross, sparse, use_s, s_k
    def retrieve(self, q, top_k=5):
        if not q or not q.strip(): return []
        q_emb = self.em.generate_embeddings([q])[0]
        if self.use_s and self.sparse and self.sparse.is_ready():
            sparse_cands = self.sparse.get_top_n(q, self.s_k)
            raw = self.vs.collection.query(query_embeddings=[q_emb.tolist()], n_results=DENSE_RETRIEVAL_CANDIDATES, include=["documents", "metadatas", "embeddings"])
            if not raw["documents"][0]: return []
            dense_cands = [{"id": raw["ids"][0][i], "content": raw["documents"][0][i], "metadata": raw["metadatas"][0][i], "emb": raw["embeddings"][0][i]} for i in range(len(raw["documents"][0]))]
            cands = self._hybrid_fusion(q_emb, sparse_cands, dense_cands)
        else:
            raw = self.vs.collection.query(query_embeddings=[q_emb.tolist()], n_results=DENSE_RETRIEVAL_CANDIDATES, include=["documents", "metadatas", "embeddings"])
            if not raw["documents"][0]: return []
            cands = [{"id": raw["ids"][0][i], "content": raw["documents"][0][i], "metadata": raw["metadatas"][0][i], "emb": raw["embeddings"][0][i]} for i in range(len(raw["documents"][0]))]
            sims = cosine_similarity([q_emb], [c["emb"] for c in cands])[0]
            for i, s in enumerate(sims): cands[i]["similarity_score"] = float(s)
        if self.cross is not None:
            # Optimize: Only rerank top 20 candidates instead of all for better performance
            cands = self._rerank_with_cross_encoder(q, cands, min(20, top_k * CROSS_ENCODER_CANDIDATE_MULTIPLIER))
        return sorted(cands, key=lambda x: x.get("similarity_score", 0.0), reverse=True)[:top_k]

    def _hybrid_fusion(self, q_emb, sparse_cands, dense_cands):
        all_cands = {}
        sparse_scores = [c.get("sparse_score", 0.0) for c in sparse_cands]
        max_sparse = max(sparse_scores) if sparse_scores else 1.0
        for rank, c in enumerate(sparse_cands):
            cid = c.get("id", sha256_text_shared(c.get("content", "")))
            norm = c.get("sparse_score", 0.0) / max(max_sparse, 1e-10)
            all_cands[cid] = {**c, "sparse_rank": rank, "sparse_score_norm": norm}
        dense_sims = cosine_similarity([q_emb], [c["emb"] for c in dense_cands])[0]
        for rank, (c, sim) in enumerate(zip(dense_cands, dense_sims)):
            cid = c.get("id", sha256_text_shared(c.get("content", "")))
            if cid not in all_cands: all_cands[cid] = {**c, "dense_rank": rank, "dense_score": float(sim)}
            else: all_cands[cid].update({"dense_rank": rank, "dense_score": float(sim)})
        for cid, c in all_cands.items():
            s_rank, d_rank = c.get("sparse_rank", DEFAULT_MISSING_RANK), c.get("dense_rank", DEFAULT_MISSING_RANK)
            rrf = (1.0 / (s_rank + RRF_RANK_CONSTANT)) + (1.0 / (d_rank + RRF_RANK_CONSTANT))
            c["similarity_score"] = float(HYBRID_FUSION_RRF_WEIGHT * rrf + HYBRID_FUSION_DENSE_WEIGHT * c.get("dense_score",0.0) + HYBRID_FUSION_SPARSE_WEIGHT * c.get("sparse_score_norm",0.0))
        return list(all_cands.values())

    def _rerank_with_cross_encoder(self, query, candidates, top_n):
        try:
            pairs = [[query, c["content"]] for c in candidates]
            scores = self.cross.predict(pairs)
            for i, s in enumerate(scores):
                # Apply sigmoid to normalize raw logits to a 0.0 - 1.0 range
                norm_s = 1.0 / (1.0 + np.exp(-float(s)))
                candidates[i]["cross_encoder_score"] = norm_s
                candidates[i]["similarity_score"] = ORIGINAL_SCORE_WEIGHT * candidates[i].get("similarity_score", 0.0) + CROSS_ENCODER_WEIGHT * norm_s
            return sorted(candidates, key=lambda x: x.get("similarity_score", 0.0), reverse=True)[:top_n]
        except Exception: return candidates

class GroqLLMRAG:
    def __init__(self, model="llama-3.1-8b-instant", key=None):
        self.llm = ChatGroq(groq_api_key=key, model_name=model, temperature=0.1)
    def generate_response(self, q, ctx, prompt_instr=""):
        p = f"""{prompt_instr}
TASK: Answer the user's question using ONLY the provided context.
CONTEXT:
{ctx}
QUESTION: {q}
INSTRUCTIONS:
1. Answer based EXCLUSIVELY on context.
2. If partial, state "Partial information available:".
3. If no info, state "Insufficient information."
4. Distinguish Mandatory (must/shall) vs Optional.
5. Professional, short paragraphs.Cite specific details & pages.
6. Output: Plain text response.
ANSWER:"""
        return self.llm.invoke([HumanMessage(content=p)]).content

    def expand_query(self, q):
        try:
            prompt = f"""Generate {MAX_QUERY_VARIATIONS} alternative phrasings of: "{q}"
Respond with ONLY the alternative questions, one per line. VARIATIONS:"""
            response = self.llm.invoke([HumanMessage(content=prompt)]).content
            variations = [l.strip() for l in response.strip().split('\n') if l.strip() and len(l.strip())>3]
            return [q] + variations[:MAX_QUERY_VARIATIONS]
        except: return [q]

    def validate_answer(self, q, answer, context):
        try:
            prompt = f"Grounded in context?\nQ: {q}\nA: {answer}\nC: {context[:1000]}\nStart with YES/NO and brief reason. VALIDATION:"
            response = self.llm.invoke([HumanMessage(content=prompt)]).content
            return response.strip().upper().startswith('YES'), response
        except: return True, "Validation error"

def assemble_context_rag(chunks, max_ctx=4096):
    parts = [f"Source: {c['metadata'].get('source_file')} (Page: {c['metadata'].get('page', 'N/A')}) [ID: {c['id'][:8]}]\n{c['content']}" for c in chunks]
    full = "\n\n".join(parts)
    if len(full) > max_ctx * CHARS_PER_TOKEN_ESTIMATE: full = full[:max_ctx * CHARS_PER_TOKEN_ESTIMATE]
    return full, chunks

def multi_query_retrieval(retriever, queries, top_k=5):
    all_res = {}
    for q in queries:
        res = retriever.retrieve(q, top_k=top_k * MULTI_QUERY_RETRIEVAL_MULTIPLIER)
        for rank, r in enumerate(res):
            id_ = r.get("id", sha256_text_shared(r["content"]))
            if id_ not in all_res: all_res[id_] = {**r, "ranks": [], "scores": []}
            all_res[id_]["ranks"].append(rank); all_res[id_]["scores"].append(r.get("similarity_score", 0.0))
    for r in all_res.values():
        rrf = sum(1.0 / (rk + RRF_RANK_CONSTANT) for rk in r["ranks"])
        avg = sum(r["scores"]) / len(r["scores"])
        r["similarity_score"] = MULTI_QUERY_RRF_WEIGHT * rrf + MULTI_QUERY_AVG_WEIGHT * avg
    return sorted(all_res.values(), key=lambda x: x.get("similarity_score", 0.0), reverse=True)[:top_k]

def extract_sentences_from_chunks(chunks):
    all_s = []
    for i, c in enumerate(chunks):
        content = c.get("content", "")
        sents = simple_sent_tokenize_shared(content)
        pos = 0
        for s in sents:
            start = content.find(s, pos)
            if start == -1: start = pos
            all_s.append({"text": s, "chunk_idx": i, "chunk_id": c.get("id", ""), "source_file": c.get("metadata", {}).get("source_file", ""), "page": c.get("metadata", {}).get("page", "N/A"), "char_start": start, "char_end": start + len(s)})
            pos = start + len(s)
    return all_s

def map_sentences_to_sources_rag(answer, used, em):
    a_sents = simple_sent_tokenize_shared(answer)
    if not a_sents or not used: return []
    try:
        s_sents = extract_sentences_from_chunks(used)
        if not s_sents: return []
        # Limit the number of source sentences to avoid OOM on large contexts
        s_sents = s_sents[:MAX_CITATION_SENTENCES]
        a_embs, s_embs = em.generate_embeddings(a_sents), em.generate_embeddings([s["text"] for s in s_sents])
        sims = cosine_similarity(a_embs, s_embs)
        cites = []
        for i, ans in enumerate(a_sents):
            idx = int(np.argmax(sims[i]))
            best_s = s_sents[idx]
            cites.append({"answer_sentence": ans, "source_sentence": best_s["text"], "source_file": best_s["source_file"], "page": best_s["page"], "chunk_id": best_s["chunk_id"], "similarity_score": float(sims[i][idx]), "char_range": f"{best_s['char_start']}-{best_s['char_end']}"})
        return cites
    except: return []

def extract_clauses_from_chunks(chunks):
    clauses = []
    combined = '|'.join(f'({p})' for p in CLAUSE_PATTERNS)
    for i, c in enumerate(chunks):
        content = c.get("content", "")
        parts = re.split(combined, content)
        curr, pos = "", 0
        for p in parts:
            if p and p.strip():
                curr += p
                if len(curr.strip()) > MIN_CLAUSE_LENGTH:
                    start = content.find(curr.strip(), pos)
                    if start == -1: start = pos
                    clauses.append({"text": curr.strip(), "chunk_idx": i, "chunk_id": c.get("id", ""), "source_file": c.get("metadata", {}).get("source_file", ""), "page": c.get("metadata", {}).get("page", "N/A"), "char_start": start, "char_end": start + len(curr.strip())})
                    pos, curr = start + len(curr.strip()), ""
        if curr.strip() and len(curr.strip()) > MIN_CLAUSE_LENGTH:
            clauses.append({"text": curr.strip(), "chunk_idx": i, "chunk_id": c.get("id", ""), "source_file": c.get("metadata", {}).get("source_file", ""), "page": c.get("metadata", {}).get("page", "N/A"), "char_start": pos, "char_end": pos + len(curr.strip())})
    return clauses

def map_answer_to_clauses_rag(answer, used, em):
    a_sents = simple_sent_tokenize_shared(answer)
    if not a_sents or not used: return []
    try:
        clauses = extract_clauses_from_chunks(used)
        if not clauses: return []
        # Limit the number of clauses to avoid OOM on large contexts
        clauses = clauses[:MAX_CITATION_SENTENCES]
        a_embs, c_embs = em.generate_embeddings(a_sents), em.generate_embeddings([c["text"] for c in clauses])
        sims = cosine_similarity(a_embs, c_embs)
        cites = []
        for i, ans in enumerate(a_sents):
            idx = int(np.argmax(sims[i]))
            best_c = clauses[idx]
            txt = best_c["text"]
            if len(txt) > MAX_CLAUSE_DISPLAY_LENGTH: txt = txt[:MAX_CLAUSE_DISPLAY_LENGTH] + "..."
            cites.append({"answer_sentence": ans, "source_clause": txt, "source_file": best_c["source_file"], "page": best_c["page"], "chunk_id": best_c["chunk_id"], "similarity_score": float(sims[i][idx]), "char_range": f"{best_c['char_start']}-{best_c['char_end']}"})
        return cites
    except: return []

# --------------------------
# Summarization Components
# --------------------------

class ConcurrencyManager:
    """Manages a global semaphore that is safe across Streamlit event loop re-runs.

    A plain asyncio.Semaphore created at module level becomes bound to the event
    loop that was current when it was first *used*.  Because Streamlit calls
    asyncio.run() on every analysis run, a new event loop is created each time
    and the old semaphore silently hangs or raises RuntimeError.  This manager
    recreates the semaphore whenever the running event loop changes.
    """
    def __init__(self, value: int):
        self.value = value
        self._sem = None
        self._loop = None

    def get(self):
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            return None  # Should not happen inside asyncio.run()
        if self._sem is None or self._loop != loop:
            self._sem = asyncio.Semaphore(self.value)
            self._loop = loop
        return self._sem


_SUM_CONCURRENCY_MGR = ConcurrencyManager(2)


class TokenBucket:
    """Token-bucket rate limiter that is safe across Streamlit event loop re-runs.

    asyncio.Lock is also event-loop-bound, so the lock is recreated whenever the
    running loop changes (same pattern as ConcurrencyManager above).
    """
    def __init__(self, rpm: int):
        self.spacing = 60.0 / max(1, rpm)
        self.last_call = 0.0
        self._lock = None

    @property
    def lock(self):
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            if self._lock is None:
                self._lock = asyncio.Lock()
            return self._lock
        if not hasattr(self, '_loop_id'):
            self._loop_id = None
        current_loop_id = id(loop)
        if self._lock is None or self._loop_id != current_loop_id:
            self._lock = asyncio.Lock()
            self._loop_id = current_loop_id
        return self._lock

    async def wait(self):
        async with self.lock:
            # Use time.monotonic() instead of asyncio.get_event_loop().time() so
            # that self.last_call stays consistent across asyncio.run() restarts.
            # loop.time() resets to near-zero on each new event loop, which caused
            # self.last_call (from the previous run) to be far in the "future"
            # relative to the new loop — resulting in forced 60-second sleeps on
            # every API call in subsequent analysis runs.
            now = time.monotonic()
            if self.last_call < now:
                self.last_call = now
            wait_time = (self.last_call + self.spacing) - now
            if wait_time > 0:
                await asyncio.sleep(min(wait_time, 60.0))
                self.last_call += self.spacing
            else:
                self.last_call = now + self.spacing + random.uniform(0.05, 0.20)

    async def report_429(self, wait_sec: float):
        """Global push-back when a 429 is received."""
        async with self.lock:
            now = time.monotonic()
            self.last_call = max(self.last_call, now + wait_sec)
            try:
                st.session_state["api_backoff_until"] = time.time() + wait_sec
            except Exception:
                pass


_sum_rate_limiters: dict = {}


def get_sum_limiter(rpm: int) -> TokenBucket:
    if rpm not in _sum_rate_limiters:
        _sum_rate_limiters[rpm] = TokenBucket(rpm)
    return _sum_rate_limiters[rpm]

def clean_json_string(txt):
    if not txt: return "{}"
    txt = txt.strip()
    if "```" in txt:
        try:
            parts = re.findall(r"```(?:json)?\s*([\s\S]*?)\s*```", txt)
            txt = parts[0].strip() if parts else re.sub(r"^```(?:json)?\s*|\s*```$", "", txt).strip()
        except: pass
    if not (txt.startswith("{") or txt.startswith("[")):
        brace_pos = txt.find("{")
        bracket_pos = txt.find("[")
        start = min((p for p in [brace_pos, bracket_pos] if p != -1), default=-1)
        if start != -1 and start < len(txt):
            end = txt.rfind("]") if txt[start] == "[" else txt.rfind("}")
            if end > start: txt = txt[start:end+1]
    txt = re.sub(r',(\s*[}\]])', r'\1', txt)
    # Replace Python literals with JSON equivalents
    txt = re.sub(r'\bTrue\b', 'true', txt)
    txt = re.sub(r'\bFalse\b', 'false', txt)
    txt = re.sub(r'\bNone\b', 'null', txt)
    # NOTE: Do NOT apply an unquoted-key regex here — it would corrupt evidence
    # strings that contain ", word:" patterns (e.g. "Section 3, clause: bidder shall"),
    # turning valid JSON into invalid JSON and causing all map results to be discarded.
    return txt

def fix_json_string_newlines(txt: str) -> str:
    """Escape unescaped control characters (newlines, tabs, carriage returns)
    inside JSON string values to prevent 'Unterminated string' parse errors."""
    result = []
    in_str = False
    escaped = False
    ESCAPES = {'\n': '\\n', '\r': '\\r', '\t': '\\t'}
    for ch in txt:
        if escaped:
            result.append(ch)
            escaped = False
        elif ch == '\\' and in_str:
            result.append(ch)
            escaped = True
        elif ch == '"':
            in_str = not in_str
            result.append(ch)
        elif in_str and ch in ESCAPES:
            result.append(ESCAPES[ch])
        else:
            result.append(ch)
    return ''.join(result)


def repair_truncated_json(txt: str) -> str:
    """Close any unterminated JSON string and open bracket/brace structures
    to produce parseable JSON when the LLM response is truncated."""
    stack = []
    in_str = False
    escaped = False
    for ch in txt:
        if escaped:
            escaped = False
            continue
        if ch == '\\' and in_str:
            escaped = True
            continue
        if ch == '"':
            in_str = not in_str
            continue
        if not in_str:
            if ch in '{[':
                stack.append(ch)
            elif ch in '}]':
                if stack:
                    stack.pop()
    closers = []
    if in_str:
        closers.append('"')
    for bracket in reversed(stack):
        closers.append('}' if bracket == '{' else ']')
    return txt + ''.join(closers)


async def call_gemini_json_sum_async(client, sys, user, model, rpm):
    """Robust Gemini JSON call using ConcurrencyManager + TokenBucket rate limiting.

    Mirrors the call_gemini_json_async pattern from Summarizationcode.py:
    - Uses ConcurrencyManager.get() so the semaphore is always bound to the
      *current* event loop (safe across multiple asyncio.run() calls in Streamlit).
    - Uses TokenBucket for proper inter-call spacing and 429 back-off.
    - Up to 20 retries with exponential back-off, matching the working pipeline
      in Summarizationcode.py.
    """
    limiter = get_sum_limiter(rpm)
    backoff = 5.0

    for attempt in range(20):
        try:
            await limiter.wait()
            sem = _SUM_CONCURRENCY_MGR.get()
            from google.genai import types
            gen_config = types.GenerateContentConfig(
                system_instruction=sys,
                response_mime_type="application/json"
            )
            if sem:
                async with sem:
                    resp = await client.aio.models.generate_content(
                        model=model, contents=user, config=gen_config
                    )
            else:
                resp = await client.aio.models.generate_content(
                    model=model, contents=user, config=gen_config
                )
            txt = clean_json_string(resp.text or "{}")
            try:
                return json.loads(txt)
            except json.JSONDecodeError:
                # Fix 1: escape raw control characters (newlines etc.) inside strings
                fixed = fix_json_string_newlines(txt)
                try:
                    return json.loads(fixed)
                except json.JSONDecodeError:
                    # Fix 2: repair truncated / unclosed JSON structures
                    try:
                        return json.loads(repair_truncated_json(fixed))
                    except json.JSONDecodeError:
                        pass
                # Fix 3: convert single-quoted keys/values to double-quoted
                txt = re.sub(r"(^|\s|{|,)\s*'([^']+)'\s*:", r'\1"\2":', txt)
                txt = re.sub(r":\s*'([^']*)'(\s*[,}\]])", r': "\1"\2', txt)
                try:
                    return json.loads(txt)
                except Exception as e:
                    return {"error": f"JSON parse error: {str(e)}", "raw": txt[:500]}
        except Exception as e:
            last_err = str(e)
            msg = last_err.lower()
            if "429" in msg or "resource_exhausted" in msg or "rate" in msg or "quota" in msg:
                wait_sec = backoff + random.random() * 2
                match = re.search(r"retrydelay':\s*'(\d+)s", msg)
                if match:
                    try:
                        wait_sec = float(match.group(1)) + 1
                    except (ValueError, AttributeError):
                        pass
                await limiter.report_429(wait_sec)
                await asyncio.sleep(wait_sec)
                backoff = min(backoff * 2.0, 120.0)
                continue
            if "400" in msg or "invalid_argument" in msg:
                return {"error": f"Invalid request: {last_err}"}
            if attempt == 19:
                return {"error": last_err}
            await asyncio.sleep(backoff + random.random())
            backoff *= 1.5
    return {"error": "max_retries_exceeded"}


def _extract_reduce_items(red: dict, s_obj: str) -> list:
    """Extract the list of findings from an intermediate reduce result dict.

    Used by _hierarchical_reduce to collect items from each batch-reduce
    response before the final consolidation pass.
    """
    if not isinstance(red, dict) or red.get("error"):
        return []
    key_map = {
        "Compliance Matrix": "matrix",
        "Risk Assessment": "risks",
        "Ambiguity Scrutiny": "ambiguities",
    }
    key = key_map.get(s_obj)
    if key:
        val = red.get(key)
        return val if isinstance(val, list) else []
    if s_obj == "Entity Dashboard":
        dashboard = red.get("dashboard")
        if not isinstance(dashboard, dict):
            return []
        items = []
        for cat, cat_items in dashboard.items():
            if isinstance(cat_items, list):
                for item in cat_items:
                    if isinstance(item, dict):
                        item_copy = dict(item)
                        item_copy.setdefault("category", cat)
                        items.append(item_copy)
        return items
    # General Summary / Overall Summary & Voice
    return red.get("key_findings") or []


async def _hierarchical_reduce(client, rs, ri, q, mapped, s_model, s_rpm, s_obj):
    """Two-phase hierarchical reduce with **parallel** batch reduces.

    Phase 1: All batches of at most REDUCE_BATCH_THRESHOLD findings are reduced
             concurrently via asyncio.gather, eliminating the sequential wait
             that caused long delays and heightened 429 exposure.
    Phase 2: All intermediate items are merged in a single final reduce call.

    Falls back to raw mapped items when every batch reduce fails so the final
    reduce still has material to work with.
    """
    try:
        if len(mapped) <= REDUCE_BATCH_THRESHOLD:
            return await call_gemini_json_sum_async(
                client, rs, f"{ri}\nQ:{q}\nD:{json.dumps(mapped)}", s_model, s_rpm
            )

        batches = [
            mapped[i:i + REDUCE_BATCH_THRESHOLD]
            for i in range(0, len(mapped), REDUCE_BATCH_THRESHOLD)
        ]
        st.write(f"  Parallel reducing {len(batches)} batch(es) ({len(mapped)} findings)...")

        # Run all batch reduces concurrently to minimize wall-clock latency.
        batch_results = await asyncio.gather(
            *[
                call_gemini_json_sum_async(
                    client, rs, f"{ri}\nQ:{q}\nD:{json.dumps(b)}", s_model, s_rpm
                )
                for b in batches
            ],
            return_exceptions=True,
        )

        intermediate_items: list = []
        for batch, br in zip(batches, batch_results):
            if isinstance(br, Exception):
                continue  # skip failed batches; the final merge still proceeds
            extracted = _extract_reduce_items(br, s_obj)
            if extracted:
                intermediate_items.extend(extracted)
            elif isinstance(br, dict) and not br.get("error"):
                # Reduce returned a valid dict but no extractable items;
                # preserve raw batch findings so nothing is silently dropped.
                intermediate_items.extend(batch)

        if not intermediate_items:
            intermediate_items = list(mapped)

        st.write(f"  Final merge of {len(intermediate_items)} consolidated items...")
        return await call_gemini_json_sum_async(
            client, rs,
            f"{ri}\nQ:{q}\nD:{json.dumps(intermediate_items)}",
            s_model, s_rpm,
        )
    except Exception as exc:
        return {"error": f"Reduce phase error: {exc}"}


def _apply_result_fallback(mapped: list, red, s_obj: str) -> dict:
    """Unified fallback: rebuild reduce result from raw mapped items when the
    reduce phase returned an error, an empty result, or a non-dict.

    Covers all four structured objectives (Compliance Matrix, Risk Assessment,
    Entity Dashboard, Ambiguity Scrutiny) and the summary objectives.
    Returns a guaranteed dict.
    """
    if not isinstance(red, dict):
        red = {}
    if not mapped:
        return red
    red.pop("error", None)

    # Normalize page (singular) → pages (list) on every mapped item once.
    for item in mapped:
        if isinstance(item, dict) and "page" in item and "pages" not in item:
            item["pages"] = [item["page"]]

    valid = [m for m in mapped if isinstance(m, dict) and not m.get("error")]

    if s_obj == "Compliance Matrix" and not red.get("matrix"):
        red["matrix"] = [n for n in (_normalize_compliance_item(m) for m in mapped) if n is not None]
    elif s_obj == "Risk Assessment" and not red.get("risks"):
        red["risks"] = valid
    elif s_obj == "Ambiguity Scrutiny" and not red.get("ambiguities"):
        red["ambiguities"] = valid
    elif s_obj == "Entity Dashboard" and (
        not red.get("dashboard")
        or not any(isinstance(v, list) and v for v in red.get("dashboard", {}).values())
    ):
        rebuilt: dict = {}
        for m in valid:
            cat = m.get("category", "General")
            rebuilt.setdefault(cat, []).append({
                "entity": m.get("entity", ""),
                "context": m.get("context", ""),
                "evidence": m.get("evidence", ""),
                "pages": m["pages"] if isinstance(m.get("pages"), list) else ([m.get("page")] if m.get("page") else []),
            })
        if rebuilt:
            red["dashboard"] = rebuilt
    return red


def _recover_evidence_in_result(mapped: list, red: dict, s_obj: str, chunks: list) -> None:
    """In-place evidence recovery for structured reduce results.

    Fills empty 'evidence' fields in the reduce output by searching:
    1. Exact / substring match against the mapped-item evidence index.
    2. Word-overlap fuzzy match (≥ 2 shared words).
    3. Best-matching raw chunk sentence as a last resort.

    Operates on all four structured objectives; no-ops for summary objectives.
    """
    if not isinstance(red, dict):
        return

    # Build a flat evidence index from all mapped items.
    ev_index: dict[str, str] = {}
    for m in mapped:
        if isinstance(m, dict) and m.get("evidence"):
            for field in ("item", "clause", "entity", "ambiguous_text"):
                val = m.get(field)
                if val:
                    ev_index[str(val).lower().strip()] = m["evidence"]

    stop_words = _EVIDENCE_STOP_WORDS

    def _find_evidence(primary: str, secondary: str = "") -> str:
        key = primary.lower().strip()
        if not key:
            return ""
        # Strategy 1: substring / exact match.
        for k, v in ev_index.items():
            if key in k or k in key:
                return v
        # Strategy 2: word-overlap fuzzy match.
        key_words = set(key.split())
        best_ev, best_overlap = "", 0
        for k, v in ev_index.items():
            overlap = len(key_words & set(k.split()))
            if overlap > best_overlap:
                best_overlap, best_ev = overlap, v
        if best_ev and best_overlap >= 2:
            return best_ev
        # Strategy 3: search raw chunks for the best-matching sentence.
        if chunks:
            search_words = {
                w for w in (set((primary + " " + secondary).lower().split()) - stop_words)
                if len(w) > 3
            }
            best_chunk, best_score = None, 0
            for chunk in chunks:
                cl = chunk["text"].lower()
                score = sum(1 for w in search_words if w in cl)
                if score > best_score:
                    best_score, best_chunk = score, chunk
            min_score = 1 if len(search_words) <= 2 else 2
            if best_chunk and best_score >= min_score:
                try:
                    sents = simple_sent_tokenize_shared(best_chunk["text"])
                    best_sent = max(
                        sents,
                        key=lambda s: sum(1 for w in search_words if w in s.lower()),
                        default="",
                    )
                    return best_sent.strip()
                except Exception:
                    pass
        return ""

    if s_obj == "Compliance Matrix":
        for entry in red.get("matrix") or []:
            if isinstance(entry, dict) and not entry.get("evidence"):
                entry["evidence"] = _find_evidence(
                    str(entry.get("item") or ""), str(entry.get("detail") or "")
                )
    elif s_obj == "Risk Assessment":
        for entry in red.get("risks") or []:
            if isinstance(entry, dict) and not entry.get("evidence"):
                entry["evidence"] = _find_evidence(
                    str(entry.get("clause") or ""), str(entry.get("reason") or "")
                )
    elif s_obj == "Entity Dashboard":
        for items in (red.get("dashboard") or {}).values():
            if isinstance(items, list):
                for entry in items:
                    if isinstance(entry, dict) and not entry.get("evidence"):
                        entry["evidence"] = _find_evidence(
                            str(entry.get("entity") or ""), str(entry.get("context") or "")
                        )
    elif s_obj == "Ambiguity Scrutiny":
        for entry in red.get("ambiguities") or []:
            if isinstance(entry, dict) and not entry.get("evidence"):
                entry["evidence"] = _find_evidence(
                    str(entry.get("ambiguous_text") or ""), str(entry.get("issue") or "")
                )


def extract_pages_sum(path):
    """
    Extract full text from every PDF page for the summarisation pipeline.
    - Block-sorted extraction (correct reading order)
    - pdfplumber table extraction with dropna + Markdown formatting
    - Text cleaning (null bytes, redundant whitespace)
    Mirrors the extract_pages() approach from Summarizationcode.py.
    """
    doc = None
    try:
        doc, p, f = fitz.open(path), [], []
        with pdfplumber.open(path) as pl:
            for i, page in enumerate(doc):
                # Block-sorted extraction preserves reading order.
                try:
                    blocks = page.get_text("blocks")
                    blocks = sorted(blocks, key=lambda b: (round(b[1], 1), round(b[0], 1)))
                    t = "\n".join(
                        b[4] for b in blocks if isinstance(b, (list, tuple)) and len(b) >= 5
                    )
                except Exception:
                    t = page.get_text()

                # Table extraction — clean empty rows/cols and emit as Markdown
                # so the LLM sees structured data instead of raw whitespace.
                if i < len(pl.pages):
                    tabs = pl.pages[i].extract_tables()
                    for tab in tabs:
                        if tab:
                            df_t = pd.DataFrame(tab).dropna(how="all").dropna(axis=1, how="all")
                            if not df_t.empty:
                                t += (
                                    f"\n\n[TABLE DATA - PAGE {i + 1}]:\n"
                                    + df_t.to_markdown(index=False)
                                    + "\n\n"
                                )

                # Normalise whitespace and remove null bytes.
                t = t.replace("\x00", " ").replace("\r\n", "\n")
                t = re.sub(r"[ \t]+", " ", t)
                t = re.sub(r"\n{3,}", "\n\n", t).strip()

                p.append(t)
                f.append(len(page.get_images()) > 0)
        return p, f
    finally:
        if doc:
            doc.close()

def semantic_chunk_pages_sum(pages, flags, max_tok=8000, overlap=5):
    """
    Chunk extracted pages into token-limited segments for the map/reduce pipeline.

    Improvements over the previous flat-sentence approach (mirrors
    Summarizationcode.py semantic_chunk_pages logic):
    - Pages are first split into paragraphs so sentence boundaries are not
      confused across unrelated blocks.
    - Table-like blocks are kept together as individual row items so tabular
      evidence is not split across chunks arbitrarily.
    - Each sentence is prefixed with [Page N] so the LLM can cite exact pages.
    - end_page is tracked in addition to start_page for multi-page chunks.
    """
    def _looks_like_table(text):
        lines = [ln for ln in text.splitlines() if ln.strip()]
        if not lines:
            return False
        numeric_lines = sum(1 for ln in lines if re.search(r'\d', ln))
        has_delims = any(
            '|' in ln or '\t' in ln or re.search(r'\s{2,}', ln) for ln in lines
        )
        return (numeric_lines / max(1, len(lines))) > 0.25 and has_delims

    flat = []
    for i, (content, has_img) in enumerate(zip(pages, flags)):
        content = content.strip()
        if not content:
            continue
        # Split each page into paragraphs first for better semantic units.
        paragraphs = re.split(r'\n\s*\n', content)
        for par in paragraphs:
            par = par.strip()
            if not par:
                continue
            if _looks_like_table(par):
                # Emit each row as its own flat item so table data is
                # never split mid-row.
                rows = [f"[Page {i + 1}] {r.strip()}" for r in par.splitlines() if r.strip()]
                for r in rows:
                    flat.append((i + 1, r, has_img))
            else:
                try:
                    sents = nltk.sent_tokenize(par) if _NLTK_READY else simple_sent_tokenize_shared(par)
                except Exception:
                    sents = simple_sent_tokenize_shared(par)
                if sents:
                    # Prefix first sentence with page number for citation tracking.
                    sents[0] = f"[Page {i + 1}] " + sents[0]
                for s in sents:
                    flat.append((i + 1, s, has_img))

    chunks, bs, bp, bf = [], [], [], []
    for pg, s, flag in flat:
        if sum(len(x) // 4 for x in bs + [s]) > max_tok and bs:
            chunks.append({
                "text": " ".join(bs).strip(),
                "start_page": min(bp),
                "end_page": max(bp),
                "has_visual": any(bf),
                "id": str(uuid.uuid4())[:8],
            })
            bs = bs[-overlap:] if overlap > 0 else []
            bp = bp[-overlap:] if overlap > 0 else []
            bf = bf[-overlap:] if overlap > 0 else []
        bs.append(s)
        bp.append(pg)
        bf.append(flag)

    if bs:
        chunks.append({
            "text": " ".join(bs).strip(),
            "start_page": min(bp),
            "end_page": max(bp),
            "has_visual": any(bf),
            "id": str(uuid.uuid4())[:8],
        })
    return [c for c in chunks if c["text"].strip()]

def get_sum_prompts(mode: str):
    """Return (map_system, map_instruction, reduce_system, reduce_instruction).

    The MAP phase uses lightweight, format-focused prompts (adapted from
    Summarizationcode.py's get_prompts) so the LLM is not over-constrained
    during the scanning pass.  Strict mandatory-evidence rules in
    get_sum_prompts_full were causing the LLM to return empty arrays for the
    four specific objectives (Compliance Matrix, Risk Assessment, Entity
    Dashboard, Ambiguity Scrutiny).

    The REDUCE phase still uses the detailed prompts from get_sum_prompts_full
    so the final consolidated output retains high quality.
    """
    if mode == "Compliance Matrix":
        ms = "Extract mandatory tender requirements (Technical, Financial, Legal)."
        mi = ('Return JSON array: [{"item": "short label", "detail": "requirement description", '
              '"evidence": "verbatim quote from text", "category": "Technical/Legal/Financial/Administrative/Safety/General", '
              '"mandatory": true, "page": 1}]')
    elif mode == "Risk Assessment":
        ms = "Flag high-risk tender clauses (Liabilities, Penalties, Termination)."
        mi = ('Return JSON array: [{"clause": "topic", "reason": "why it is risky", '
              '"evidence": "verbatim quote from text", "risk_level": "High/Medium/Low", '
              '"risk_type": "...", "impact": "...", "page": 1}]')
    elif mode == "Entity Dashboard":
        ms = "Extract critical tender metadata: Organizations, People, Locations, Dates, Deadlines, Financials, Contacts, Technicals."
        mi = ('Return JSON array: [{"category": "Organization/People/Location/Date/Deadline/Financial/Contact/Technical", '
              '"entity": "...", "context": "brief context", "evidence": "verbatim quote from text", "page": 1}]')
    elif mode == "Ambiguity Scrutiny":
        ms = "Identify ambiguous, conflicting, or vague clauses that need clarification from the authority."
        mi = ('Return JSON array: [{"ambiguous_text": "vague term or clause", '
              '"ambiguity_type": "vague/conflicting/missing/unclear", "issue": "description of the problem", '
              '"evidence": "verbatim quote from text", "suggested_query": "formal question for bidder meet", '
              '"severity": "High/Medium/Low", "page": 1}]')
    elif mode == "Overall Summary & Voice":
        ms = "You are a comprehensive tender analyst."
        mi = ('Return JSON array: [{"domain": "Scope/Timeline/Financial/Technical/Legal", '
              '"topic": "topic title", "detail": "full description", '
              '"evidence": "verbatim quote from text", "importance": "High/Medium/Low", "page": 1}]')
    else:
        ms = "Extract key details, scope, deliverables, or any informative sections relevant to the query."
        mi = ('Return JSON array: [{"finding": "topic title", "detail": "important details or exact quote", '
              '"type": "...", "importance": "High/Medium/Low", "page": 1}]')
    _map_sys, _map_ins, rs, ri = get_sum_prompts_full(mode)
    return ms, mi, rs, ri

def get_sum_prompts_full(mode: str):
    if mode == "Compliance Matrix":
        map_system = "You are a compliance requirements extractor analyzing tender documents. Extract specific, actionable requirements."
        map_instruction = """Extract ALL compliance requirements from the provided Context Data.
Rules:
- Identify ALL obligation and requirement language, including: must/shall/will/should/required/mandatory/needs to/has to/is to/are to/obligated to/expected to.
- Also extract technical specifications, performance criteria, standards references, and conditional clauses even if no explicit modal verb is used.
Each chunk in the Context Data is separated by "---" and has a header line "ID:... P:..." followed by "Text: <content>".
Format: JSON array [{item, detail, evidence, category, mandatory, page}]
- item: a short label for the requirement (e.g. "Submission Deadline", "Technical Specification").
- detail: full description of the requirement.
- evidence: copy the relevant sentence or clause from the "Text:" section of the chunk (the text after "Text:" up to the next "---" separator) that contains the requirement. If no clear evidence sentence is found, set this field to an empty string "".
- category: classify as one of Technical, Legal, Financial, Administrative, Safety, or General.
- mandatory: true if the language is binding (must/shall/required/mandatory), false otherwise.
- page: use the P: value from the chunk header (the "ID:... P:..." line) that contains the requirement.
Rule: Include EVERY compliance requirement found in the text. Do NOT omit any requirement, even if evidence is unavailable -- set evidence to "" in that case. Every item must have a valid page number from the chunk header. If the chunk contains no compliance requirements at all, return an empty array []."""
        reduce_system = "You are consolidating compliance requirements into a unified matrix."
        reduce_instruction = """Consolidate the findings in D: into a unique compliance matrix. Deduplicate similar requirements.
CRITICAL: Preserve the exact 'evidence' text from each finding verbatim — do NOT modify, summarize, or omit the evidence values. Include ALL matrix items from the source findings — do NOT exclude any item, even if its evidence field is empty.
IMPORTANT: For each matrix item, collect and list all unique page numbers from the source findings in the 'pages' array. Do NOT leave pages empty.
Format: JSON {matrix: [{item, detail, evidence, category, mandatory, pages:[]}], total_requirements, summary: {mandatory_count, optional_count, categories: {}}}"""
    elif mode == "Risk Assessment":
        map_system = "You are a risk analyst identifying risks, liabilities, and concerns."
        map_instruction = """Extract ALL potential risks from the provided Context Data. Look for: Penalties, liabilities, tight deadlines, ambiguous clauses, legal exposure.
Each chunk in the Context Data is separated by "---" and has a header line "ID:... P:..." followed by "Text: <content>".
Format: JSON array [{clause, reason, evidence, risk_level, risk_type, impact, page}]
- evidence: copy the relevant sentence or clause from the "Text:" section of the chunk (the text after "Text:" up to the next "---" separator) that contains the risk. If no clear verbatim sentence can be extracted, set this field to an empty string "".
- page: use the P: value from the chunk header that contains the risk.
RULE: Include every risk found in the text. Do not omit any risk, even if direct evidence cannot be extracted verbatim — set evidence to "" in that case rather than omitting the item. Every item must have a valid page number from the chunk header."""
        reduce_system = "You are consolidating risk assessments."
        reduce_instruction = """Consolidate finding D: into unique risks. Prioritize High/Critical.
IMPORTANT: For each risk item, collect and list all unique page numbers from the source findings into the 'pages' array. Do NOT leave pages empty.
CRITICAL: Preserve the exact 'evidence' text from each finding verbatim — do NOT modify, summarize, or omit the evidence values. Include ALL risk items from the source findings — do NOT exclude any item, even if its evidence field is empty.
Format: JSON {risks: [{clause, reason, evidence, risk_level, risk_type, impact, pages:[]}], risk_summary: {total_risks, critical_count, high_count, medium_count, low_count, by_type: {}}}"""
    elif mode == "Entity Dashboard":
        map_system = "You are extracting key entities and metadata."
        map_instruction = """Extract Organizations, People, Locations, Dates, Deadlines, Financials, Contacts, Technicals from the provided Context Data.
Each chunk in the Context Data is separated by "---" and has a header line "ID:... P:..." followed by "Text: <content>".
Format: JSON array [{category, entity, context, evidence, page}]
- evidence: MANDATORY - copy the EXACT verbatim phrase/clause from the "Text:" section of the chunk (text after "Text:" up to the next "---") that contains the entity. Do NOT paraphrase or leave blank.
- page: use the P: value from the chunk header that contains the entity.
RULE: Every item in the output array MUST have a non-empty evidence field and a valid page number using the P: value from the chunk header."""
        reduce_system = "You are compiling entities into a dashboard."
        reduce_instruction = """Consolidate finding D: into dashboard categories. Remove duplicates. Preserve evidence and page numbers for each entity.
CRITICAL: Preserve the exact 'evidence' text from each finding verbatim — do NOT modify, summarize, or omit the evidence values. Every entity MUST have a non-empty evidence field — any entity lacking evidence is invalid and must be excluded.
IMPORTANT: For each entity, collect all unique page numbers from the source findings into the 'pages' array. Do NOT leave pages empty.
Format: JSON {dashboard: {Organizations:[{entity, context, evidence, pages:[]}], People:[{entity, context, evidence, pages:[]}], Locations:[{entity, context, evidence, pages:[]}], Dates:[{entity, context, evidence, pages:[]}], Deadlines:[{entity, context, evidence, pages:[]}], Financials:[{entity, context, evidence, pages:[]}], Contacts:[{entity, context, evidence, pages:[]}], Technicals:[{entity, context, evidence, pages:[]}]}, entity_count: {}}"""
    elif mode == "Ambiguity Scrutiny":
        map_system = "You are an expert analyst identifying ambiguous, vague, or contradictory language in tender documents."
        map_instruction = """Extract instances of vague terms ("reasonable", "appropriate"), unclear requirements, contradictions between sections, and missing details from the provided Context Data.
Each chunk in the Context Data is separated by "---" and has a header line "ID:... P:..." followed by "Text: <content>".
Format: JSON array [{ambiguous_text, ambiguity_type, issue, evidence, suggested_query, severity, page}]
- ambiguous_text: the specific vague or unclear term/phrase identified (e.g., "reasonable timeframe").
- evidence: copy the EXACT verbatim sentence or clause from the "Text:" section of the chunk (text after "Text:" up to the next "---") that contains the ambiguous text. If no clear verbatim sentence can be extracted, set this field to an empty string "".
- suggested_query: a formal question for the authority to clarify the point.
- page: use the P: value from the chunk header that contains the ambiguous text.
RULE: Include every ambiguity found in the text. Do not omit any ambiguity, even if direct evidence cannot be extracted verbatim — set evidence to "" in that case rather than omitting the item. Every item must have a valid page number from the chunk header."""
        reduce_system = "You are consolidating ambiguity findings into a formal pre-bid query report."
        reduce_instruction = """Consolidate finding D: into a comprehensive Scrutiny Report. Refine the suggested_query for each finding.
CRITICAL: Preserve the exact 'evidence' text from each finding verbatim — do NOT modify, summarize, or omit the evidence values.
IMPORTANT: For each ambiguity, collect and list all unique page numbers from the source findings into the 'pages' array. Do NOT leave pages empty.
Format: JSON {ambiguities: [{ambiguous_text, ambiguity_type, issue, evidence, suggested_query, severity, pages:[], recommendation}], summary, overall_assessment: "Brief assessment."}"""
    elif mode == "Overall Summary & Voice":
        map_system = "You are a senior analyst extracting high-density data for a 15-minute executive briefing."
        map_instruction = """Extract ALL critical information across these domains from the provided Context Data:
1. Project Scope & Objectives
2. Timelines & Milestones
3. Financials (Budget, Penalties, Payment Terms)
4. Technical Features (Mechanical, Electrical, Automation)
5. Legal & Liability
Each chunk in the Context Data is separated by "---" and has a header line "ID:... P:..." followed by "Text: <content>".
Format: JSON array [{domain, topic, detail, evidence, importance, page}]
- evidence: MANDATORY - copy the EXACT verbatim sentence from the "Text:" section of the chunk (text after "Text:" up to the next "---"). Do NOT paraphrase or leave blank.
- page: use the P: value from the chunk header that contains the information.
RULE: Every item in the output array MUST have a non-empty evidence field and a valid page number from the chunk header."""
        reduce_system = "You are a briefing expert synthesizing a massive, detailed tender overview."
        reduce_instruction = """Synthesize D: into an extremely detailed, long-form narrative summary designed to be read as a 10-15 minute briefing.
You MUST include dedicated sections for:
- Executive Overview & Project Goals
- Scope of Work & Deliverables
- Master Timeline & Critical Milestones
- Financial Breakdown & Commercial Terms
- Technical Domain: Mechanical Features
- Technical Domain: Electrical Systems
- Technical Domain: Automation & Control 
- Risk & Liability Assessment
CRITICAL: Preserve the exact 'evidence' text from each finding verbatim — do NOT modify, summarize, or omit the evidence values. For each key_finding, collect and list all unique page numbers from the source findings in the 'pages' array. Do NOT leave pages empty.
Format: JSON {summary: "500+ words overview", key_findings: [{finding, detail, domain, pages:[]}], audio_script: "A 2000+ word detailed narrative script for the audio briefing", citations: []}"""
    else:
        map_system = "You are extracting key findings from tender documents."
        map_instruction = """Extract topics, critical requirements, facts, figures from the provided Context Data.
Each chunk in the Context Data is separated by "---" and has a header line "ID:... P:..." followed by "Text: <content>".
Format: JSON array [{finding, detail, type, importance, page}]
- page: use the P: value from the chunk header that contains the finding.
RULE: Every item in the output array MUST have a valid page number from the chunk header."""
        reduce_system = "You are synthesizing findings."
        reduce_instruction = """Synthesize D: into a coherent narrative. IMPORTANT: For each key finding, you MUST collect and list all unique page numbers from the source snippets in a 'pages' array.
Format: JSON {summary: "2-3 paragraphs", key_findings: [{finding, detail, importance, pages:[]}], citations: [], finding_stats: {}}"""
    return (map_system, map_instruction, reduce_system, reduce_instruction)

# Maximum character length used when truncating a detail string to produce a fallback item label.
_MAX_FALLBACK_ITEM_LENGTH = 80


def _chunk_page_label(chunk: dict) -> str:
    """Return a page-range label for a chunk, e.g. '5' or '5-8'.

    Used to build the 'P:...' field in map-phase payloads so the LLM can
    correctly attribute evidence that falls on any page within a multi-page
    chunk, not just the first page.
    """
    start = chunk.get('start_page', '')
    end = chunk.get('end_page', start)
    if end and end != start:
        return f"{start}-{end}"
    return str(start)

# Priority-ordered key names to look for when the LLM wraps the findings array in a dict
# (e.g. {"risks": [...]} instead of the bare [...]).  Using a priority list prevents
# accidentally picking a different list-valued key (e.g. "citations" before "risks").
_MODE_LIST_KEYS: dict[str, list[str]] = {
    "Compliance Matrix": ["matrix", "requirements", "compliance_requirements", "items", "findings"],
    "Risk Assessment": ["risks", "risk_items", "risk_findings", "findings", "items"],
    "Ambiguity Scrutiny": ["ambiguities", "ambiguity_findings", "findings", "items"],
    "Entity Dashboard": ["entities", "findings", "items"],
    "Overall Summary & Voice": ["key_findings", "findings", "items"],
    "General Summary": ["key_findings", "findings", "items"],
}


def _unwrap_batch_result(b, mode: str) -> list:
    """Extract a list of finding dicts from a single batch LLM response.

    Handles three formats returned by Gemini:
    - A bare JSON array  → returned directly (filtered for valid dicts).
    - A JSON object wrapping the array → tries mode-specific keys first, then
      falls back to the first list-valued key, then treats the whole dict as a
      single finding.
    - None / error dict → returns an empty list.
    """
    if b is None:
        return []
    if isinstance(b, list):
        return [item for item in b if isinstance(item, dict) and not item.get("error")]
    if isinstance(b, dict) and not b.get("error"):
        # Try mode-specific keys first for reliable extraction.
        # Skip keys whose list is empty or contains only error dicts so we
        # can fall through to a later key that may have real findings.
        for key in _MODE_LIST_KEYS.get(mode, []):
            if isinstance(b.get(key), list) and b[key]:
                items = [item for item in b[key] if isinstance(item, dict) and not item.get("error")]
                if items:
                    return items
        # Fall back to the first non-empty list-valued key found in the dict
        for k in b:
            if isinstance(b[k], list) and b[k]:
                items = [item for item in b[k] if isinstance(item, dict) and not item.get("error")]
                if items:
                    return items
        # Last resort: treat the whole dict as a single finding,
        # but ONLY when it has no list-valued keys.
        # A dict that still has list-valued keys at this point (e.g.
        # {"matrix": [], "requirements": []}) is a container response whose
        # lists were all empty or all-error — it is NOT a direct finding.
        # Returning it as a finding would populate `mapped` with useless
        # meta-dicts and prevent the recovery extraction from triggering.
        if b and not any(isinstance(v, list) for v in b.values()):
            return [b]
    return []

def _resolve_mandatory_val(val) -> bool:
    """Normalise various LLM representations of mandatory/optional to a bool.

    The LLM may return booleans, strings such as 'true'/'yes'/'mandatory',
    or integers.  All are accepted.
    """
    if isinstance(val, bool):
        return val
    if isinstance(val, str):
        return val.strip().lower() in ("true", "yes", "mandatory", "must", "shall", "1", "required")
    if isinstance(val, (int, float)):
        return bool(val)
    return False

def _normalize_compliance_item(entry: dict) -> dict | None:
    """Return a canonical compliance-matrix dict from a raw LLM output entry.

    Accepts any of the common key variants the LLM might emit
    (e.g. 'requirement' instead of 'item', 'description' instead of
    'detail', a string 'true' for mandatory, etc.).
    Returns None when the entry has no usable text at all.
    """
    if not isinstance(entry, dict) or "error" in entry:
        return None
    # Use str() to guard against the LLM returning numeric or other non-string
    # values for text fields (e.g. "item": 1).  Without str(), the chained
    # `or ""` expression evaluates to the truthy non-string value and the
    # subsequent .strip() call raises AttributeError, crashing the reduce phase.
    item_text = str(
        entry.get("item") or entry.get("requirement") or entry.get("compliance_item")
        or entry.get("name") or entry.get("clause") or entry.get("topic")
        or entry.get("requirement_name") or ""
    ).strip()
    detail_text = str(
        entry.get("detail") or entry.get("description") or entry.get("details")
        or entry.get("specifics") or entry.get("content") or ""
    ).strip()
    if not item_text and not detail_text:
        return None
    mandatory_raw = (
        entry.get("mandatory") if entry.get("mandatory") is not None
        else entry.get("type") or entry.get("requirement_type") or False
    )
    return {
        "item": item_text or detail_text[:_MAX_FALLBACK_ITEM_LENGTH],
        "detail": detail_text,
        "evidence": str(entry.get("evidence") or ""),
        "category": str(entry.get("category") or entry.get("type") or "General"),
        "mandatory": _resolve_mandatory_val(mandatory_raw),
        "pages": (
            entry.get("pages") if isinstance(entry.get("pages"), list)
            else ([entry["page"]] if entry.get("page") else [])
        ),
    }


def compute_confidence_score_sum(mapped, reduced, q):
    conf = {"snippet_coverage": 0.0, "result_coherence": 0.0, "information_density": 0.0, "citation_confidence": 0.0, "overall_confidence": 0.0}
    if not mapped: return conf
    # 1. Snippet Coverage: 3+ snippets for full score (less punishing than 5)
    conf["snippet_coverage"] = min(1.0, len(mapped) / 3.0)
    
    # 2. Result Coherence: 1+ significant field for full score
    coh = 0.0
    if isinstance(reduced, dict) and "error" not in reduced:
        fields = ["summary", "matrix", "risks", "dashboard", "ambiguities", "key_findings"]
        present = sum(1 for f in fields if reduced.get(f))
        coh = min(1.0, present / 1.0)
    conf["result_coherence"] = coh
    
    # 3. Information Density: Better sweet spot (3x to 40x query length)
    q_tokens = tokenize_simple_shared(q); res_text = ""
    if isinstance(reduced, dict):
        for f in ["summary", "matrix", "risks", "dashboard", "ambiguities", "key_findings"]:
            val = reduced.get(f)
            if val: 
                if isinstance(val, (list, dict)): res_text += json.dumps(val)
                else: res_text += str(val)
    res_tokens = tokenize_simple_shared(res_text)
    if q_tokens and res_tokens:
        ratio = len(res_tokens) / len(q_tokens)
        if ratio < 3: conf["information_density"] = ratio / 3.0
        elif ratio <= 40: conf["information_density"] = 1.0
        else: conf["information_density"] = max(0.5, 1.0 - (ratio - 40) / 100.0)
    
    # 4. Citation Confidence: 1 citation per 3 snippets
    cites = []
    if isinstance(reduced, dict):
        # Extract citations from all potential fields
        cites = reduced.get("citations") or []
        for f in ["matrix", "risks", "ambiguities", "key_findings"]:
            items = reduced.get(f)
            if isinstance(items, list):
                for item in items:
                    if isinstance(item, dict):
                        c = item.get("pages") or item.get("page")
                        if c: cites.extend(c if isinstance(c, list) else [c])
        # Extract citations from Entity Dashboard nested structure
        dashboard = reduced.get("dashboard")
        if isinstance(dashboard, dict):
            for cat_items in dashboard.values():
                if isinstance(cat_items, list):
                    for item in cat_items:
                        if isinstance(item, dict):
                            c = item.get("pages") or item.get("page")
                            if c: cites.extend(c if isinstance(c, list) else [c])
    unique_cites = len(set(c for c in cites if c is not None))
    expected = max(1, len(mapped) / 3)
    conf["citation_confidence"] = min(1.0, unique_cites / expected)
    
    # Weights: Coverage (25%), Coherence (25%), Density (20%), Grounding (30%)
    conf["overall_confidence"] = (0.25 * conf["snippet_coverage"] + 0.25 * conf["result_coherence"] + 
                                  0.20 * conf["information_density"] + 0.30 * conf["citation_confidence"])
    return conf

def render_page_level_citations_sum(chunks, pages):
    if not chunks or not pages: return
    st.markdown("---")
    with st.expander("📋 Citation Tracking: Page-Level Extraction", expanded=False):
        st.markdown("**Source pages and text excerpts:**")
        page_chunks = {}
        for c in chunks:
            # Index the chunk under every page it spans so multi-page chunks
            # appear under each of their source pages, not just start_page.
            start = c.get('start_page', 0)
            end = c.get('end_page', start)
            for p in range(start, end + 1):
                if p in pages:
                    if p not in page_chunks: page_chunks[p] = []
                    page_chunks[p].append(c)
        for p_num in sorted(page_chunks.keys()):
            st.markdown(f"#### 📄 Page {p_num}")
            for idx, c in enumerate(page_chunks[p_num]):
                txt = c.get('text', '')
                disp = txt[:MAX_CLAUSE_DISPLAY_LENGTH] + "..." if len(txt) > MAX_CLAUSE_DISPLAY_LENGTH else txt
                st.markdown(f"**Cit {idx+1}** (ID: `{c.get('id','NA')}`) - {disp}")
                if len(txt) > MAX_CLAUSE_DISPLAY_LENGTH: 
                    with st.expander("View Full Text"): st.text(txt)
            st.markdown("---")

def render_citation_preview_sum(doc, cites):
    if not cites or not doc: return
    st.markdown("### 📄 Context Preview")
    tabs = st.tabs([f"Page {c['page']}" for c in cites])
    for i, c in enumerate(cites):
        with tabs[i]:
            try:
                p = doc.load_page(c['page'] - 1)
                pix = p.get_pixmap(matrix=fitz.Matrix(ZOOM_LEVEL, ZOOM_LEVEL))
                st.image(pix.tobytes("png"), use_container_width=True)
            except: st.error(f"Failed to load Page {c['page']}")

def convert_result_to_dataframe(result, objective):
    if not isinstance(result, dict): return None
    try:
        df = None
        if objective == "Compliance Matrix" and "matrix" in result: df = pd.DataFrame(result["matrix"])
        elif objective == "Risk Assessment" and "risks" in result: df = pd.DataFrame(result["risks"])
        elif objective == "Entity Dashboard" and "dashboard" in result:
            entities = []
            for cat, items in result["dashboard"].items():
                if isinstance(items, list):
                    for it in items:
                        if isinstance(it, dict): it_c = it.copy(); it_c["category"] = cat; entities.append(it_c)
            df = pd.DataFrame(entities) if entities else pd.DataFrame(columns=["category", "entity", "context", "evidence", "pages"])
        elif objective == "Ambiguity Scrutiny" and "ambiguities" in result: df = pd.DataFrame(result["ambiguities"])
        elif objective == "General Summary":
            for k in result.keys():
                if isinstance(result[k], list) and result[k]: df = pd.DataFrame(result[k]); break
        if df is not None and "pages" in df.columns:
            df["pages"] = df["pages"].apply(lambda x: ", ".join(map(str, x)) if isinstance(x, list) else str(x))
        return df
    except: return None

def export_to_excel(df, query, objective):
    if df is None or df.empty: return None
    try:
        from io import BytesIO
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Results')
            ws = writer.sheets['Results']
            for col in ws.columns:
                max_l = max([len(str(cell.value)) for cell in col])
                ws.column_dimensions[col[0].column_letter].width = min(max_l + 2, 50)
        return output.getvalue()
    except: return None

def export_to_word(df, query, objective, result=None):
    try:
        from io import BytesIO
        output = BytesIO()
        doc = Document()
        doc.add_heading(f'{objective} Analysis', 0)
        doc.add_paragraph(f'Query: {query}')
        
        if objective in ("General Summary", "Overall Summary & Voice") and result:
            # Narrative format for General Summary and Overall Summary & Voice
            if result.get("summary"):
                doc.add_heading('Executive Summary', level=1)
                doc.add_paragraph(result["summary"])
            
            if result.get("key_findings"):
                doc.add_heading('Key Findings', level=1)
                for item in result["key_findings"]:
                    heading = item.get("finding", "Finding")
                    doc.add_heading(heading, level=2)
                    doc.add_paragraph(item.get("detail", "N/A"))
                    if item.get("domain"):
                        doc.add_paragraph(f"Domain: {item['domain']}")
                    if item.get("pages"):
                        p_str = ", ".join(map(str, item["pages"]))
                        doc.add_paragraph(f"Source Pages: {p_str}").italic = True
            
            if result.get("overall_assessment"):
                doc.add_heading('Overall Assessment', level=1)
                doc.add_paragraph(result["overall_assessment"])
            
            if objective == "Overall Summary & Voice" and result.get("audio_script"):
                doc.add_heading('Audio Script', level=1)
                doc.add_paragraph(result["audio_script"])
        else:
            # Default table format for other objectives
            if df is None or df.empty: return None
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Light Grid Accent 1'
            # Header
            for i, col in enumerate(df.columns): table.rows[0].cells[i].text = str(col)
            # Data
            for _, row in df.iterrows():
                cells = table.add_row().cells
                for i, v in enumerate(row): cells[i].text = str(v) if v is not None else ''
        
        doc.save(output)
        return output.getvalue()
    except Exception as e:
        print(f"Word export error: {e}")
        return None

def sanitize_filename(text, max_length=30):
    for c in '<>:"/\\|?*': text = text.replace(c, '_')
    return text[:max_length].strip().replace(' ', '_')

def generate_audio_briefing(text):
    if not text: return None
    try:
        from io import BytesIO
        tts = gTTS(text=text, lang='en', slow=False)
        audio_io = BytesIO()
        tts.write_to_fp(audio_io)
        return audio_io.getvalue()
    except Exception as e:
        print(f"TTS Error: {e}")
        return None

# --------------------------
# UI Helper Components
# --------------------------
def render_confidence_gauge(value):
    color = "#EF4444" if value < 0.3 else "#F59E0B" if value < 0.7 else "#10B981"
    # Convert 0-1.0 to 0-180 degrees
    angle = int(value * 180)
    
    st.markdown(f"""
        <div class="gauge-container">
            <div style="font-weight: 600; color: #374151; margin-bottom: 10px;">Analysis Confidence</div>
            <div class="gauge-wrap">
                <div class="gauge-bg"></div>
                <div class="gauge-fill" style="background: conic-gradient(from -90deg at 50% 50%, {color} 0deg, {color} {angle}deg, #e5e7eb {angle}deg);"></div>
                <div class="gauge-inner" style="color: {color};">
                    {value:.1%}
                </div>
            </div>
            <div class="legend-container">
                <div class="legend-item"><div class="legend-color" style="background: #EF4444;"></div> < 30% (Low)</div>
                <div class="legend-item"><div class="legend-color" style="background: #F59E0B;"></div> 30-70% (Mid)</div>
                <div class="legend-item"><div class="legend-color" style="background: #10B981;"></div> > 70% (High)</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

def render_metric_cards(mapped_len, total_time, confidence):
    color = "#EF4444" if confidence < 0.3 else "#F59E0B" if confidence < 0.7 else "#10B981"
    angle = int(confidence * 180)
    
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"""<div class="metric-card">
            <p style="color: #6b7280; font-size: 0.875rem; margin-bottom: 0;">Context Coverage</p>
            <p style="font-size: 1.5rem; font-weight: 700; margin: 0;">{mapped_len} Chunks</p>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="metric-card">
            <p style="color: #6b7280; font-size: 0.875rem; margin-bottom: 0;">Total Processing Time</p>
            <p style="font-size: 1.5rem; font-weight: 700; margin: 0;">{total_time:.2f}s</p>
        </div>""", unsafe_allow_html=True)
    with c3:
        # Integrated Dial Gauge in the 3rd column
        st.markdown(f"""
            <div class="metric-card" style="padding: 0.5rem;">
                <div class="gauge-container" style="width: 100%; margin: 0 auto; transform: scale(0.85); transform-origin: top center;">
                    <p style="color: #6b7280; font-size: 0.875rem; margin-bottom: 5px;">Confidence Score</p>
                    <div class="gauge-wrap" style="width: 140px; height: 70px;">
                        <div class="gauge-bg" style="width: 140px; height: 140px;"></div>
                        <div class="gauge-fill" style="width: 140px; height: 140px; background: conic-gradient(from -90deg at 50% 50%, {color} 0deg, {color} {angle}deg, #e5e7eb {angle}deg);"></div>
                        <div class="gauge-inner" style="width: 110px; height: 55px; left: 15px; font-size: 1.1rem; color: {color};">
                            {confidence:.1%}
                        </div>
                    </div>
                    <div class="legend-container" style="gap: 8px; margin-top: 5px; font-size: 0.65rem;">
                        <div class="legend-item"><div class="legend-color" style="width: 8px; height: 8px; background: #EF4444;"></div> 0-30%</div>
                        <div class="legend-item"><div class="legend-color" style="width: 8px; height: 8px; background: #F59E0B;"></div> 30-70%</div>
                        <div class="legend-item"><div class="legend-color" style="width: 8px; height: 8px; background: #10B981;"></div> 70-100%</div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

# --------------------------
# Main Program logic
# --------------------------
with st.sidebar:
    st.markdown("""<div style="text-align: center; margin-bottom: 1rem;">
        <h2 style="color: #1e3a8a;">Tender Analyzer Pro</h2>
        <p style="color: #64748b;">Enterprise Search & Analysis</p>
    </div>""", unsafe_allow_html=True)
    choice = st.radio("Navigation", ["⚡ Simple QA (RAG)", "🧠 Context Understanding"], label_visibility="collapsed")
    
    st.markdown("---")
    st.subheader("⚙️ Global Settings")
    
    # Combined sidebar for both modes with expanders
    if choice == "⚡ Simple QA (RAG)":
        ga_key = st.text_input("Groq API Key", type="password")
        with st.expander("Advanced RAG Tuning"):
            emb_m = st.text_input("Embedding model", "all-MiniLM-L6-v2")
            crs_m = st.text_input("Cross-Encoder", "cross-encoder/ms-marco-MiniLM-L-6-v2")
            u_crs = st.checkbox("Use Cross-Encoder", True)
            u_spa = st.checkbox("Use Sparse BM25", True)
            chunk_method = st.selectbox("Chunking Method", ["Recursive", "Hybrid", "Fixed-size"])
            c_size = st.number_input("Chunk size", 800)
            c_over = st.number_input("Overlap", 128)
            max_tk = st.number_input("Context Tokens", 4096)
            top_k = st.number_input("Top K UI", 3)
            supp_th = st.slider("Support threshold", 0.0, 1.0, 0.45)
    else:
        gem_key = st.text_input("Google API Key", type="password")
        s_obj = st.selectbox("Analysis Objective", ["General Summary", "Overall Summary & Voice", "Compliance Matrix", "Risk Assessment", "Entity Dashboard", "Ambiguity Scrutiny"])
        with st.expander("Advanced Summary Tuning"):
            s_model = st.text_input("Model", value="gemini-2.5-flash")
            s_rpm = st.number_input("Target RPM", 1, 60, 4)
            s_batch = st.number_input("Batch size", 5, 50, 10)
            s_max_tk = st.number_input("Max tokens/chunk", 500, 32000, 8000)
            s_over = st.number_input("Overlap sentences", 0, 20, 5)

# --------------------------
# RAG Execution
# --------------------------
if choice == "⚡ Simple QA (RAG)":
    st.markdown('<div class="app-header"><h1>⚡ Tender QA (RAG Mode)</h1><p>Ask specific questions about your tender documents with precise citations.</p></div>', unsafe_allow_html=True)
    files = st.file_uploader("Upload Tender PDFs", type="pdf", accept_multiple_files=True)

    # _get_em and _get_vs are defined at module level so @st.cache_resource
    # correctly identifies and reuses them across Streamlit reruns.

    if files:
        if st.button("🚀 Index Documents", use_container_width=True):
            with st.status("Indexing Documents...", expanded=True) as status:
                try:
                    em = _get_em(emb_m)
                    vs = _get_vs()

                    st.write("Extracting Text & Metadata...")
                    all_pages = []
                    for f in files:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as t:
                            t.write(f.getbuffer()); t.flush()
                            t_path = t.name
                        try:
                            pages = process_pdf(t_path)
                            # Override source_file with original filename so citations
                            # reference the user-uploaded file name, not the temp path.
                            for p in pages:
                                p.metadata["source_file"] = f.name
                            if pages:
                                all_pages.extend(pages)
                        finally:
                            try: os.unlink(t_path)
                            except OSError: pass

                    if not all_pages:
                        st.error("No text could be extracted from the uploaded PDFs. Please check the files and try again.")
                        status.update(label="Indexing Failed", state="error", expanded=True)
                    else:
                        chunks = split_documents(all_pages, chunk_size=int(c_size), chunk_overlap=int(c_over), method=chunk_method)
                        del all_pages

                        if not chunks:
                            st.error("No chunks generated. Please try different chunking settings.")
                            status.update(label="Indexing Failed", state="error", expanded=True)
                        else:
                            st.write(f"Generating Embeddings for {len(chunks)} chunks...")
                            if len(chunks) > MAX_CHUNKS_WARNING_THRESHOLD:
                                st.warning(f"⚠️ Large document set detected ({len(chunks)} chunks). "
                                           "Embeddings are generated in batches to limit memory usage.")
                            texts = [d.page_content for d in chunks]
                            emb_batches = []
                            for i in range(0, len(texts), EMBEDDING_BATCH_SIZE):
                                emb = em.generate_embeddings(texts[i:i + EMBEDDING_BATCH_SIZE])
                                emb_batches.append(emb)
                            all_emb = np.vstack(emb_batches) if emb_batches else np.zeros((0, DEFAULT_EMBEDDING_DIM), dtype=np.float32)
                            del emb_batches, texts

                            if all_emb.size == 0:
                                st.error("Embeddings generation failed.")
                                status.update(label="Indexing Failed", state="error", expanded=True)
                            else:
                                vs.add_documents(chunks, all_emb)
                                del all_emb

                                st.session_state["rvs_v2"], st.session_state["sem_v2"] = vs, em

                                if u_spa and BM25_AVAILABLE:
                                    st.write("Building Sparse BM25 Index...")
                                    bm = SparseBM25IndexRAG(); bm.build(chunks)
                                    st.session_state["rbm_v2"] = bm

                                del chunks

                                st.write("Success!")
                                status.update(label="Index Complete!", state="complete", expanded=False)
                except Exception as e:
                    st.error(f"Indexing failed. Please check your documents and try again. (Details: {e})")
                    status.update(label="Indexing Failed", state="error", expanded=True)

    if "rvs_v2" in st.session_state:
        # Question input area
        q = st.text_input("Enter your question about the uploaded tenders:")
        
        # Move display options to main area for better visibility
        with st.expander("🛠️ Result Display Options", expanded=True):
            col1, col2, col3, col4 = st.columns(4)
            with col1: use_qe = st.checkbox("Query Expansion", True, help="Generate multiple queries for better retrieval")
            with col2: use_av = st.checkbox("Answer Validation", True, help="LLM-based grounding check")
            with col3: show_cl = st.checkbox("Clause Citations", True, help="Map answer to specific document clauses")
            with col4: show_sc = st.checkbox("Sentence Citations", True, help="Map answer to specific source sentences")

        if q:
            if not ga_key: st.error("Groq API Key Mandatory"); st.stop()
            with st.status("Processing Request...", expanded=True) as status:
                t_start = time.time()
                cr = load_cross_encoder_cached(crs_m) if u_crs and CROSS_ENCODER_AVAILABLE else None
                llm = GroqLLMRAG(key=ga_key)
                ret = RAGRetrieverRAG(st.session_state["rvs_v2"], st.session_state["sem_v2"], cr, sparse=st.session_state.get("rbm_v2"), use_s=u_spa)
                
                if use_qe:
                    st.write("Expanding Query...")
                    queries = llm.expand_query(q)
                    res = multi_query_retrieval(ret, queries, top_k=top_k)
                else: res = ret.retrieve(q, top_k)
                
                t_ret = time.time() - t_start
                st.write("Generating Answer...")
                ctx, used = assemble_context_rag(res, max_ctx=max_tk)
                ans = llm.generate_response(q, ctx)
                t_gen = time.time() - (t_start + t_ret)
                t_total = time.time() - t_start
                status.update(label="Answer Ready!", state="complete", expanded=False)

            # Display Results in V2 Dashboard Style
            st.markdown('<div class="stCard">', unsafe_allow_html=True)
            st.subheader("💡 Analysis Result")
            st.write(strip_provenance(ans))
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Confidence & Metrics
            supp = compute_mean_support_score_shared(res)
            render_metric_cards(len(res), t_total, supp)
            
            tab1, tab2, tab3 = st.tabs(["📚 Evidence Chunks", "📋 Precise Citations", "✅ Validation"])
            with tab1:
                for i, c in enumerate(res):
                    with st.expander(f"Chunk {i+1} | Score: {c.get('similarity_score',0.0):.3f} | Page: {c['metadata'].get('page')}"):
                        st.text(c['content'])
            with tab2:
                if show_cl:
                    st.markdown("#### 📑 Clause-Level Citations")
                    c_map = map_answer_to_clauses_rag(strip_provenance(ans), used, st.session_state["sem_v2"])
                    if not c_map: st.info("No clause-level matches found.")
                    for m in c_map: st.caption(f"**{m['answer_sentence']}** → *{m['source_file']} P{m['page']}*"); st.write(f"> {m['source_clause']}")
                
                if show_sc:
                    st.markdown("#### 📝 Sentence-Level Citations")
                    s_map = map_sentences_to_sources_rag(strip_provenance(ans), used, st.session_state["sem_v2"])
                    if not s_map: st.info("No sentence-level matches found.")
                    for m in s_map: st.caption(f"**{m['answer_sentence']}** → *{m['source_file']} P{m['page']}*"); st.write(f"> {m['source_sentence']}")
            with tab3:
                if use_av:
                    is_v, msg = llm.validate_answer(q, ans, ctx)
                    if is_v:
                        st.success(msg)
                    else:
                        st.warning(msg)

# --------------------------
# Summarization Execution
# --------------------------
else:
    st.markdown('<div class="app-header"><h1>🧠 Contextual Understanding</h1><p>Deep logical analysis and extraction across massive tender volumes.</p></div>', unsafe_allow_html=True)
    f_sum = st.file_uploader("Upload PDFs for Analysis", type="pdf", accept_multiple_files=True)
    q_sum = st.text_area("Analysis Queries (one per line):", placeholder="e.g., What are the total financial liabilities?")
    
    final_res = []  # initialise so display code below is safe even when button hasn't been clicked
    if st.button("🔍 Perform Deep Analysis", use_container_width=True):
        if not f_sum or not gem_key or not q_sum.strip(): st.error("Missing Files/Key/Query"); st.stop()
        
        with st.status("Analyzing Documents...", expanded=True) as status:
            client = genai.Client(api_key=gem_key)
            all_p, all_f, temp_files = [], [], []
            for f in f_sum:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as t:
                    temp_files.append(t.name); t.write(f.getbuffer()); t.flush()
                    p, fl = extract_pages_sum(t.name); all_p.extend(p); all_f.extend(fl)
            st.write(f"Chunking {len(all_p)} pages...")
            chunks = semantic_chunk_pages_sum(all_p, all_f, max_tok=s_max_tk, overlap=s_over)
            # Free raw page text — no longer needed once chunks are built.
            del all_p, all_f
            
            async def run_sum_v2():
                ms, mi, rs, ri = get_sum_prompts(s_obj)
                res = []
                batches = [chunks[i:i+s_batch] for i in range(0, len(chunks), s_batch)]
                # Pre-build all chunk payloads once (reused for retries / recovery).
                chunk_payloads = [
                    "\n\n".join(
                        f"---\nID:{c['id']} P:{_chunk_page_label(c)}\nText: {c['text']}"
                        for c in b
                    )
                    for b in batches
                ]
                for q in q_sum.splitlines():
                    if not q.strip():
                        continue
                    t0 = time.time()
                    st.write(f"Mapping query: {q}...")

                    # ----------------------------------------------------------
                    # PARALLEL MAP PHASE
                    # All chunk batches are sent to the LLM concurrently.
                    # return_exceptions=True is critical: it prevents a single
                    # 429 (rate-limit) or API error from cancelling all other
                    # concurrent tasks, so the map phase always collects as many
                    # findings as possible even when some batches fail.
                    # ----------------------------------------------------------
                    mapped_batches = await asyncio.gather(
                        *[
                            call_gemini_json_sum_async(
                                client, ms, f"{mi}\nQ:{q}\nContext Data:\n{p}", s_model, s_rpm
                            )
                            for p in chunk_payloads
                        ],
                        return_exceptions=True,
                    )
                    t_m = time.time()

                    mapped: list = []
                    failed_indices: list = []
                    for bi, b in enumerate(mapped_batches):
                        if isinstance(b, Exception):
                            failed_indices.append(bi)
                        else:
                            mapped.extend(_unwrap_batch_result(b, s_obj))
                    del mapped_batches

                    # Retry failed batches sequentially — important for Compliance
                    # Matrix where every page range must be covered.
                    if failed_indices:
                        st.write(f"Retrying {len(failed_indices)} failed batch(es)...")
                        for fi in failed_indices:
                            retry = await call_gemini_json_sum_async(
                                client, ms,
                                f"{mi}\nQ:{q}\nContext Data:\n{chunk_payloads[fi]}",
                                s_model, s_rpm,
                            )
                            mapped.extend(_unwrap_batch_result(retry, s_obj))

                    # Emergency recovery: if all map batches returned nothing,
                    # attempt sequential extraction one batch at a time.
                    if not mapped and chunks:
                        st.write("Recovery extraction in progress...")
                        for p in chunk_payloads:
                            items = _unwrap_batch_result(
                                await call_gemini_json_sum_async(
                                    client, ms, f"{mi}\nQ:{q}\nContext Data:\n{p}", s_model, s_rpm
                                ),
                                s_obj,
                            )
                            if items:
                                mapped.extend(items)
                                if s_obj != "Compliance Matrix":
                                    break  # one successful batch is enough for most modes

                    # ----------------------------------------------------------
                    # HIERARCHICAL REDUCE PHASE (parallel batch reduces)
                    # ----------------------------------------------------------
                    st.write(f"Reducing {len(mapped)} findings...")
                    red = await _hierarchical_reduce(client, rs, ri, q, mapped, s_model, s_rpm, s_obj)

                    # Unified fallback: rebuild from mapped items if reduce failed.
                    red = _apply_result_fallback(mapped, red, s_obj)

                    # Normalize Compliance Matrix item keys to canonical names.
                    if s_obj == "Compliance Matrix" and isinstance(red, dict) and red.get("matrix"):
                        red["matrix"] = [
                            n for n in (_normalize_compliance_item(e) for e in red["matrix"])
                            if n is not None
                        ]

                    # Unified evidence recovery from mapped items and raw chunks.
                    _recover_evidence_in_result(mapped, red, s_obj, chunks)

                    # Safety guard: ensure red is always a dict.
                    if not isinstance(red, dict):
                        red = (
                            {"matrix": [n for n in (_normalize_compliance_item(i) for i in red if isinstance(i, dict)) if n is not None]}
                            if isinstance(red, list) and s_obj == "Compliance Matrix"
                            else {}
                        )

                    res.append({
                        "query": q, "result": red,
                        "map_t": t_m - t0, "red_t": time.time() - t_m,
                        "mapped": mapped, "total_t": time.time() - t0,
                        "chunk_count": len(chunks),
                    })
                return res
            # Use a fresh event loop for every analysis run to avoid stale asyncio
            # state between successive objective completions without a page reload.
            # asyncio.run() is NOT used here because in some Streamlit/Python
            # configurations the script may run inside an already-running event
            # loop, causing RuntimeError.  Creating an explicit new loop guarantees
            # a clean slate regardless of Streamlit's internal async state.
            _analysis_loop = asyncio.new_event_loop()
            _analysis_exc = None
            try:
                asyncio.set_event_loop(_analysis_loop)
                final_res = _analysis_loop.run_until_complete(run_sum_v2())
            except Exception as _e:
                _analysis_exc = _e
                final_res = []
            finally:
                try:
                    _analysis_loop.close()
                except Exception:
                    pass
                asyncio.set_event_loop(asyncio.new_event_loop())
            if _analysis_exc is not None:
                status.update(label="Analysis Failed!", state="error", expanded=False)
                st.error(f"Analysis failed: {_analysis_exc}. Please try again — no page reload needed.")
                st.stop()
            status.update(label="Analysis Complete!", state="complete", expanded=False)

        for r in final_res:
            st.markdown(f"### 📋 Analysis Goal: {r['query']}")
            if isinstance(r["result"], dict) and "error" in r["result"]: st.error(r["result"]["error"]); continue
            
            conf = compute_confidence_score_sum(r['mapped'], r['result'], r['query'])
            render_metric_cards(r.get('chunk_count', len(r['mapped'])), r['total_t'], conf['overall_confidence'])

            with st.expander("🔍 Confidence Score Breakdown", expanded=False):
                sc_col1, sc_col2, sc_col3, sc_col4 = st.columns(4)
                score_details = [
                    ("📄 Snippet Coverage",    conf['snippet_coverage'],    "Weight: 25%", "Number of context chunks mapped to the query"),
                    ("🧩 Result Coherence",    conf['result_coherence'],    "Weight: 25%", "Completeness of structured AI output fields"),
                    ("📊 Information Density", conf['information_density'], "Weight: 20%", "Richness of response relative to query length"),
                    ("📌 Citation Confidence", conf['citation_confidence'], "Weight: 30%", "Unique page citations grounding the result"),
                ]
                for col, (label, score, weight, desc) in zip([sc_col1, sc_col2, sc_col3, sc_col4], score_details):
                    sc = "#EF4444" if score < 0.3 else "#F59E0B" if score < 0.7 else "#10B981"
                    with col:
                        st.markdown(f"""
                        <div style="background:#f9fafb;border:1px solid #e5e7eb;border-radius:8px;padding:0.75rem;text-align:center;">
                            <div style="font-size:0.72rem;color:#9ca3af;margin-bottom:3px;">{weight}</div>
                            <div style="font-size:0.82rem;font-weight:600;color:#374151;margin-bottom:6px;">{label}</div>
                            <div style="font-size:1.5rem;font-weight:700;color:{sc};">{score:.1%}</div>
                            <div style="font-size:0.7rem;color:#9ca3af;margin-top:5px;">{desc}</div>
                        </div>""", unsafe_allow_html=True)

            # Action Tabs for V2
            tabs_list = ["📊 Data Visualization", "🖼️ Context Preview", "📥 Export & Raw"]
            if s_obj == "Overall Summary & Voice":
                tabs_list.insert(1, "🎤 Audio Briefing")
            
            atabs = st.tabs(tabs_list)
            
            # Dynamically handle tab assignment
            visual_tab = atabs[0]
            if s_obj == "Overall Summary & Voice":
                audio_tab = atabs[1]
                context_tab = atabs[2]
                export_tab = atabs[3]
            else:
                context_tab = atabs[1]
                export_tab = atabs[2]
            
            with visual_tab:
                # Dashboard logic
                if s_obj == "Entity Dashboard" and "dashboard" in r["result"]:
                    dash = r["result"]["dashboard"]
                    # Category summary
                    cat_counts = {cat: len(items) for cat, items in dash.items() if isinstance(items, list) and items}
                    if cat_counts:
                        st.markdown("#### 📊 Category Summary")
                        cc_cols = st.columns(min(len(cat_counts), 4))
                        for i, (cat, cnt) in enumerate(cat_counts.items()):
                            with cc_cols[i % len(cc_cols)]:
                                st.metric(cat, cnt)
                        st.markdown("---")
                    for cat, items in dash.items():
                        if items and isinstance(items, list):
                            st.markdown(f"**{cat}** ({len(items)})")
                            for idx, ent in enumerate(items):
                                if not isinstance(ent, dict): continue
                                with st.container():
                                    st.markdown(f"**{idx+1}. {ent.get('entity', 'Entity')}**")
                                    st.write(f"**Context:** {ent.get('context', 'N/A')}")
                                    if ent.get('evidence'):
                                        st.markdown(f"**Evidence:** *\"{ent['evidence']}\"*")
                                    else:
                                        st.warning("**Evidence:** Not captured")
                                    st.caption(f"Pages: {ent.get('pages', [])}")
                                    st.markdown("---")
                elif s_obj == "Risk Assessment" and "risks" in r["result"]:
                    risks = r["result"]["risks"]
                    valid_risks = [rk for rk in risks if isinstance(rk, dict)]
                    # Type/Level summary
                    if valid_risks:
                        st.markdown("#### 📊 Risk Summary by Type & Level")
                        _LEVEL_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
                        type_counts = collections.Counter(rk.get('risk_type', 'Unknown') for rk in valid_risks)
                        level_counts = collections.Counter(rk.get('risk_level', 'Unknown') for rk in valid_risks)
                        rs_col1, rs_col2 = st.columns(2)
                        with rs_col1:
                            st.markdown("**By Type**")
                            for t, cnt in sorted(type_counts.items(), key=lambda x: -x[1]):
                                st.markdown(f"- {t}: **{cnt}**")
                        with rs_col2:
                            st.markdown("**By Level**")
                            for lvl, cnt in sorted(level_counts.items(), key=lambda x: _LEVEL_ORDER.get(x[0], 99)):
                                st.markdown(f"- {lvl}: **{cnt}**")
                        st.markdown("---")
                    for idx, risk in enumerate(valid_risks):
                        with st.container():
                            st.markdown(f"**{idx+1}. {risk.get('risk_type', 'Risk')}** — `{risk.get('risk_level', 'Unknown')}` severity")
                            st.write(f"**Clause:** {risk.get('clause', 'N/A')}")
                            st.write(f"**Reason:** {risk.get('reason', 'N/A')}")
                            if risk.get('evidence'):
                                st.markdown(f"**Evidence:** *\"{risk['evidence']}\"*")
                            else:
                                st.warning("**Evidence:** Not captured")
                            st.write(f"**Impact:** {risk.get('impact', 'N/A')}")
                            st.caption(f"Pages: {risk.get('pages', [])}")
                            st.markdown("---")
                    risk_summary = r["result"].get("risk_summary")
                    if risk_summary:
                        st.info(f"**Risk Summary:** Total: {risk_summary.get('total_risks', 0)} | Critical: {risk_summary.get('critical_count', 0)} | High: {risk_summary.get('high_count', 0)} | Medium: {risk_summary.get('medium_count', 0)} | Low: {risk_summary.get('low_count', 0)}")
                elif s_obj == "Ambiguity Scrutiny" and "ambiguities" in r["result"]:
                    ambs = r["result"]["ambiguities"]
                    valid_ambs = [a for a in ambs if isinstance(a, dict)]
                    # Type/Severity summary
                    if valid_ambs:
                        st.markdown("#### 📊 Ambiguity Summary by Type & Severity")
                        _LEVEL_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
                        atype_counts = collections.Counter(a.get('ambiguity_type', 'Unknown') for a in valid_ambs)
                        asev_counts = collections.Counter(a.get('severity', 'Medium') for a in valid_ambs)
                        as_col1, as_col2 = st.columns(2)
                        with as_col1:
                            st.markdown("**By Type**")
                            for t, cnt in sorted(atype_counts.items(), key=lambda x: -x[1]):
                                st.markdown(f"- {t}: **{cnt}**")
                        with as_col2:
                            st.markdown("**By Severity**")
                            for sev, cnt in sorted(asev_counts.items(), key=lambda x: _LEVEL_ORDER.get(x[0], 99)):
                                st.markdown(f"- {sev}: **{cnt}**")
                        st.markdown("---")
                    for idx, a in enumerate(valid_ambs):
                        with st.container():
                            st.markdown(f"**{idx+1}. {a.get('ambiguity_type', 'Ambiguity')}**")
                            st.write(f"**Issue:** {a.get('issue', 'N/A')}")
                            if a.get('evidence'):
                                ev = a['evidence']
                                st.markdown(f"**Evidence:** *\"{ev}\"*")
                            else:
                                st.warning("**Evidence:** Not captured")
                            st.warning(f"**Suggested Query:** {a.get('suggested_query', 'N/A')}")
                            if a.get('recommendation'): st.info(f"**Recommendation:** {a.get('recommendation')}")
                            st.caption(f"Ref: {a.get('ambiguous_text', 'Section')} | Severity: {a.get('severity', 'Medium')} | Pages: {a.get('pages', [])}")
                            st.markdown("---")
                elif s_obj == "Compliance Matrix" and "matrix" in r["result"]:
                    matrix_items = r["result"]["matrix"]
                    valid_items = [item for item in matrix_items if isinstance(item, dict)]
                    # Category summary
                    if valid_items:
                        st.markdown("#### 📊 Requirements Summary by Category")
                        cat_counts = collections.Counter(item.get('category', 'Uncategorized') for item in valid_items)
                        total_reqs = len(valid_items)
                        mandatory_count = sum(1 for item in valid_items if _resolve_mandatory_val(item.get("mandatory")))
                        optional_count = total_reqs - mandatory_count
                        cm_col1, cm_col2 = st.columns(2)
                        with cm_col1:
                            st.markdown("**By Category**")
                            for cat, cnt in sorted(cat_counts.items(), key=lambda x: -x[1]):
                                st.markdown(f"- {cat}: **{cnt}**")
                        with cm_col2:
                            st.markdown("**By Type**")
                            st.markdown(f"- 🔴 Mandatory: **{mandatory_count}**")
                            st.markdown(f"- 🟡 Optional: **{optional_count}**")
                            st.markdown(f"- Total: **{total_reqs}**")
                        st.markdown("---")
                    for idx, item in enumerate(valid_items):
                        with st.container():
                            mandatory_badge = "🔴 Mandatory" if _resolve_mandatory_val(item.get("mandatory")) else "🟡 Optional"
                            st.markdown(f"**{idx+1}. {item.get('item', 'Requirement')}** — {mandatory_badge} | Category: `{item.get('category', 'N/A')}`")
                            st.write(f"**Detail:** {item.get('detail', 'N/A')}")
                            if item.get('evidence'):
                                st.markdown(f"**Evidence:** *\"{item['evidence']}\"*")
                            else:
                                st.warning("**Evidence:** Not captured")
                            st.caption(f"Pages: {item.get('pages', [])}")
                            st.markdown("---")
                else:
                    df = convert_result_to_dataframe(r["result"], s_obj)
                    if df is not None: st.dataframe(df, use_container_width=True)
                    # Type/domain summary for General Summary and Overall Summary & Voice
                    key_findings = r["result"].get("key_findings") or []
                    valid_kf = [kf for kf in key_findings if isinstance(kf, dict)]
                    if valid_kf:
                        st.markdown("#### 📊 Findings Summary")
                        domain_counts = collections.Counter(kf.get('domain') or kf.get('type') or 'General' for kf in valid_kf)
                        kf_cols = st.columns(min(len(domain_counts), 4))
                        for i, (domain, cnt) in enumerate(sorted(domain_counts.items(), key=lambda x: -x[1])):
                            with kf_cols[i % len(kf_cols)]:
                                st.markdown(f"- {domain}: **{cnt}**")
                        st.caption(f"Total findings: **{len(valid_kf)}**")
                        st.markdown("---")
                    if "summary" in r["result"]: st.info(r["result"]["summary"])
                    if "overall_assessment" in r["result"]: st.success(f"**Overall Assessment:** {r['result']['overall_assessment']}")

            if s_obj == "Overall Summary & Voice":
                with audio_tab:
                    st.markdown("### 🎤 AI Audio Briefing")
                    script = r["result"].get("audio_script") or r["result"].get("summary")
                    if script:
                        with st.spinner("Generating High-Quality Audio Briefing..."):
                            audio_bytes = generate_audio_briefing(script)
                            if audio_bytes:
                                st.audio(audio_bytes, format="audio/mp3")
                                st.download_button("📥 Download Audio Briefing", audio_bytes, f"{sanitize_filename(r['query'])}_briefing.mp3", "audio/mp3")
                            else:
                                st.error("Failed to generate audio. Please check your internet connection/API settings.")
                    else:
                        st.warning("No script found to generate audio.")

            with context_tab:
                pgs = []
                def _f(o):
                    if isinstance(o, dict):
                        for k, v in o.items():
                            if k.lower() in ["page", "pages", "citations"]:
                                if isinstance(v, list):
                                    for x in v:
                                        try: pgs.append(int(x))
                                        except: pass
                                else:
                                    try: pgs.append(int(v))
                                    except: pass
                            else:
                                _f(v)
                    elif isinstance(o, list):
                        for x in o: _f(x)
                _f(r['result'])
                if pgs:
                    p_set = sorted(set(pgs))
                    render_page_level_citations_sum(chunks, p_set)
                    if f_sum:
                        pdf_doc = fitz.open(stream=f_sum[0].getvalue(), filetype="pdf")
                        render_citation_preview_sum(pdf_doc, [{"page": p} for p in p_set])
                        pdf_doc.close()

            with export_tab:
                df = convert_result_to_dataframe(r["result"], s_obj)
                excel_data = export_to_excel(df, r['query'], s_obj) if df is not None else None
                word_data = export_to_word(df, r['query'], s_obj, result=r['result'])
                if excel_data is not None or word_data is not None:
                    col1, col2 = st.columns(2)
                    with col1:
                        if excel_data is not None:
                            st.download_button("📥 Excel", excel_data, f"{sanitize_filename(r['query'])}.xlsx", use_container_width=True)
                    with col2:
                        if word_data is not None:
                            st.download_button("📥 Word", word_data, f"{sanitize_filename(r['query'])}.docx", use_container_width=True)
                st.subheader("Raw AI Output")
                st.json(r["result"])

            st.markdown("---")
            
        for tf in temp_files:
            try: os.unlink(tf)
            except: pass
        # Free large in-memory objects once display is complete to reduce
        # Streamlit's working-set memory before the next analysis run.
        if 'chunks' in locals():
            del chunks
        if 'final_res' in locals():
            del final_res
        gc.collect()
