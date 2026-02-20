   # MergedRAGSummarizer.py

# 1:1 functional merge of RAG and Summarization logic

# Restructured as top-down for maximum Streamlit stability.



import os

import re

import time

import json

import uuid

import hashlib

import importlib

import pickle

import tempfile

import asyncio

import random

from pathlib import Path

from typing import List, Dict, Any, Optional, Tuple



import streamlit as st

import numpy as np

import pandas as pd

import chromadb

import fitz  # PyMuPDF

import nltk

from google import genai

from docx import Document

import pdfplumber

from sentence_transformers import SentenceTransformer

from sklearn.metrics.pairwise import cosine_similarity



# --------------------------

# Page Configuration (Must be first)

# --------------------------

st.set_page_config(page_title="RAG & Summarizer", layout="wide")



# --------------------------

# Optional & Heavy Imports

# --------------------------

try:

    from sentence_transformers import CrossEncoder

    CROSS_ENCODER_AVAILABLE = True

except Exception:

    CrossEncoder = None

    CROSS_ENCODER_AVAILABLE = False



try:

    from langchain_groq import ChatGroq

    from langchain.schema import HumanMessage

    LANGCHAIN_AVAILABLE = True

except Exception:

    ChatGroq = HumanMessage = None

    LANGCHAIN_AVAILABLE = False



try:

    from bert_score import score as bert_score_fn

    BERTSCORE_AVAILABLE = True

except Exception:

    bert_score_fn = None

    BERTSCORE_AVAILABLE = False



try:

    from rank_bm25 import BM25Okapi

    BM25_AVAILABLE = True

except Exception:

    BM25Okapi = None

    BM25_AVAILABLE = False



# --------------------------

# Constants

# --------------------------

MAX_FILE_SIZE_MB = 50

MAX_QUERY_LENGTH = 5000

DEFAULT_RATE_LIMIT_SLEEP = 60.0  # seconds per request at 1 RPM

ZOOM_LEVEL = 1.5

MAX_CHUNK_TOKENS = 4096

DEFAULT_EMBEDDING_DIM = 384

CHARS_PER_TOKEN_ESTIMATE = 4



# RAG Retrieval Constants

RRF_RANK_CONSTANT = 60  # Constant k for reciprocal rank fusion: score = 1/(rank + k)

HYBRID_FUSION_RRF_WEIGHT = 0.5  # Weight for RRF score in hybrid fusion

HYBRID_FUSION_DENSE_WEIGHT = 0.3  # Weight for dense similarity score

HYBRID_FUSION_SPARSE_WEIGHT = 0.2  # Weight for sparse (BM25) score

CROSS_ENCODER_WEIGHT = 0.6  # Weight for cross-encoder score in reranking

ORIGINAL_SCORE_WEIGHT = 0.4  # Weight for original similarity score in reranking

MULTI_QUERY_RRF_WEIGHT = 0.6  # Weight for RRF in multi-query fusion

MULTI_QUERY_AVG_WEIGHT = 0.4  # Weight for average score in multi-query fusion



# Citation Constants

MIN_CLAUSE_LENGTH = 50  # Minimum character length for a valid clause

ANSWER_VALIDATION_CONTEXT_CHARS = 2000  # Max context chars for answer validation

MAX_CLAUSE_DISPLAY_LENGTH = 200  # Max chars to display for clause text in citations

MAX_CHUNK_PREVIEW_LENGTH = 300  # Max chars to display for chunk preview in UI



# Retrieval parameters

DENSE_RETRIEVAL_CANDIDATES = 100  # Number of candidates to retrieve using dense embeddings

DEFAULT_MISSING_RANK = 1000  # Penalty rank for candidates missing from one retrieval method

MAX_QUERY_VARIATIONS = 2  # Maximum number of query variations to generate (in addition to original)

MULTI_QUERY_RETRIEVAL_MULTIPLIER = 2  # Multiplier for retrieving more results per query

CROSS_ENCODER_CANDIDATE_MULTIPLIER = 3  # Multiplier for candidates before cross-encoder reranking



# Clause extraction patterns for tender documents

CLAUSE_PATTERNS = [

    r'\d+\.\d+(?:\.\d+)?',  # Section numbers like 1.1, 1.2.3

    r'[A-Z][A-Za-z\s]+:',  # Headers like "Payment:", "PAYMENT:", "Payment Terms:"

    r'\([a-z]\)',  # Sub-clauses like (a), (b)

    r'[•\-]\s+',  # Bullet points

]





# --------------------------

# Initialization

# --------------------------

@st.cache_resource

def setup_nltk_shared():

    try:

        nltk.download('punkt', quiet=True)

        nltk.download('punkt_tab', quiet=True)

        return True

    except Exception as e:

        st.warning(f"Failed to download NLTK data: {e}")

        return False



_NLTK_READY = setup_nltk_shared()



# Shared Utils

def sha256_text_shared(text: str) -> str:

    return hashlib.sha256((text or "").strip().encode("utf-8")).hexdigest()



def sha256_file_shared(path: str) -> str:

    h = hashlib.sha256()

    with open(path, "rb") as f:

        for chunk in iter(lambda: f.read(8192), b""): h.update(chunk)

    return h.hexdigest()



def normalize_text_shared(s: str) -> str:

    if s is None: return ""

    s = s.lower()

    s = re.sub(r"[^a-z0-9\s]", " ", s)

    s = re.sub(r"\s+", " ", s).strip()

    return s



def tokenize_simple_shared(s: str) -> List[str]:

    return re.findall(r"\w+", (s or "").lower())



_SENTENCE_RE = re.compile(r'(?<=[.!?])\s+')

def simple_sent_tokenize_shared(text: str) -> List[str]:

    sents = [s.strip() for s in _SENTENCE_RE.split(text or "") if s.strip()]

    return sents if sents else ([text.strip()] if (text or "").strip() else [])



_PROV_SPLIT_RE = re.compile(r"\n\s*(\[\d+\]\s*Source:|→\s*Source:|Source:|Source\s*[:\-])", flags=re.IGNORECASE)

def strip_provenance(text: str) -> str:

    if not text: return ""

    parts = _PROV_SPLIT_RE.split(text)

    cleaned = parts[0] if parts else text

    cleaned = re.sub(r"\(chunk:?[^\)]*\)", "", cleaned, flags=re.IGNORECASE)

    cleaned = re.sub(r"[\[\(]\s*\d+\s*[\]\)]", "", cleaned)

    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    return cleaned



def compute_mean_support_score_shared(results: List[Dict[str, Any]]) -> float:

    if not results: return 0.0

    scores = [float(r.get("similarity_score", 0.0)) for r in results]

    return float(np.max(scores)) if scores else 0.0



# --------------------------

# PDF Loading with PyMuPDF

# --------------------------

class Document:

    """Simple Document class to replace langchain Document"""

    def __init__(self, page_content: str, metadata: dict = None):

        self.page_content = page_content

        self.metadata = metadata or {}



def load_pdf_with_pymupdf(pdf_path: str) -> List[Document]:

    """Load PDF using PyMuPDF (fitz) and return list of Documents"""

    docs = []

    try:

        pdf_document = fitz.open(pdf_path)

        for page_num in range(len(pdf_document)):

            page = pdf_document[page_num]

            text = page.get_text()

            if text.strip():  # Only add pages with content

                doc = Document(

                    page_content=text,

                    metadata={"page": page_num + 1}

                )

                docs.append(doc)

        pdf_document.close()

    except Exception as e:

        st.error(f"Error loading PDF with PyMuPDF: {e}")

    return docs



# --------------------------

# Chunking Methods

# --------------------------

def chunk_recursive_character(docs: List[Document], chunk_size: int = 800, chunk_overlap: int = 128) -> List[Document]:

    """Recursive Character Text Splitter - splits by characters with overlap"""

    chunks = []



    for doc in docs:

        text = doc.page_content

        metadata = doc.metadata.copy()



        # Split text into chunks

        start = 0

        while start < len(text):

            end = start + chunk_size

            chunk_text = text[start:end]



            if chunk_text.strip():

                chunks.append(Document(

                    page_content=chunk_text,

                    metadata=metadata

                ))



            start += chunk_size - chunk_overlap

            if start >= len(text):

                break



    return chunks



def chunk_hybrid_token_semantic(docs: List[Document], chunk_size: int = 800, chunk_overlap: int = 128) -> List[Document]:

    """Hybrid chunking combining token-based limits with semantic boundaries (sentences)"""

    chunks = []



    for doc in docs:

        text = doc.page_content

        metadata = doc.metadata.copy()



        # Split into sentences

        try:

            if _NLTK_READY:

                sentences = nltk.sent_tokenize(text)

            else:

                sentences = simple_sent_tokenize_shared(text)

        except:

            sentences = simple_sent_tokenize_shared(text)



        current_chunk = []

        current_length = 0



        for sentence in sentences:

            sentence_length = len(sentence)



            # If adding this sentence exceeds chunk_size, save current chunk

            if current_length + sentence_length > chunk_size and current_chunk:

                chunk_text = " ".join(current_chunk)

                if chunk_text.strip():

                    chunks.append(Document(

                        page_content=chunk_text,

                        metadata=metadata

                    ))



                # Start new chunk with overlap

                # Keep last few sentences for overlap

                overlap_sentences = []

                overlap_length = 0



                for s in reversed(current_chunk):

                    if overlap_length + len(s) <= chunk_overlap:

                        overlap_sentences.insert(0, s)

                        overlap_length += len(s)

                    else:

                        break



                current_chunk = overlap_sentences + [sentence]

                current_length = overlap_length + sentence_length

            else:

                current_chunk.append(sentence)

                current_length += sentence_length



        # Add remaining chunk

        if current_chunk:

            chunk_text = " ".join(current_chunk)

            if chunk_text.strip():

                chunks.append(Document(

                    page_content=chunk_text,

                    metadata=metadata

                ))



    return chunks



def chunk_fixed_context_window(docs: List[Document], window_size: int = 800) -> List[Document]:

    """Fixed Context Window - splits text into fixed-size non-overlapping windows"""

    chunks = []



    for doc in docs:

        text = doc.page_content

        metadata = doc.metadata.copy()



        # Split text into fixed windows

        for i in range(0, len(text), window_size):

            chunk_text = text[i:i + window_size]



            if chunk_text.strip():

                chunks.append(Document(

                    page_content=chunk_text,

                    metadata=metadata

                ))



    return chunks



def apply_chunking_method(docs: List[Document], method: str, chunk_size: int = 800, chunk_overlap: int = 128) -> List[Document]:

    """Apply selected chunking method to documents"""

    if method == "Recursive Character Splitter":

        return chunk_recursive_character(docs, chunk_size, chunk_overlap)

    elif method == "Hybrid (Token+Semantic)":

        return chunk_hybrid_token_semantic(docs, chunk_size, chunk_overlap)

    elif method == "Fixed Context Window":

        return chunk_fixed_context_window(docs, chunk_size)

    else:

        # Default to recursive character splitter

        return chunk_recursive_character(docs, chunk_size, chunk_overlap)



# --------------------------

# RAG Components

# --------------------------

class EmbeddingManagerRAG:

    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):

        self.model = SentenceTransformer(model_name)

    def generate_embeddings(self, texts: List[str]) -> np.ndarray:

        return self.model.encode(texts, show_progress_bar=False)



class VectorStoreRAG:

    def __init__(self, collection_name: str = "pdf_documents", path: str = "./vector_store"):

        self.client = chromadb.PersistentClient(path=path)

        self.collection = self.client.get_or_create_collection(name=collection_name)

    def add_documents(self, docs, embs):

        ids = [f"{d.metadata.get('file_hash','na')}_{sha256_text_shared(d.page_content)}" for d in docs]

        self.collection.add(ids=ids, documents=[d.page_content for d in docs], metadatas=[d.metadata for d in docs], embeddings=embs.tolist())



class SparseBM25IndexRAG:

    def __init__(self): self.bm25, self.docs = None, []

    def build(self, docs, embs):

        self.docs = [{"id": sha256_text_shared(d.page_content), "content": d.page_content, "metadata": d.metadata, "emb": e} for d, e in zip(docs, embs)]

        self.bm25 = BM25Okapi([tokenize_simple_shared(d.page_content) for d in docs])

    def is_ready(self): return self.bm25 is not None

    def get_top_n(self, q, n):

        if not self.bm25: return []

        scores = self.bm25.get_scores(tokenize_simple_shared(q))

        idx = np.argsort(scores)[::-1][:n]

        return [dict(self.docs[i], sparse_score=float(scores[i])) for i in idx]



class RAGRetrieverRAG:

    def __init__(self, vs, em, cross=None, sparse=None, use_s=False, s_k=200):

        self.vs, self.em, self.cross, self.sparse, self.use_s, self.s_k = vs, em, cross, sparse, use_s, s_k



    def retrieve(self, q, top_k=5):

        if not q or not q.strip():

            st.warning("Empty query provided. Please enter a valid question.")

            return []



        q_emb = self.em.generate_embeddings([q])[0]



        # Hybrid retrieval: combine dense and sparse if both available

        if self.use_s and self.sparse and self.sparse.is_ready():

            # Get sparse candidates

            sparse_cands = self.sparse.get_top_n(q, self.s_k)



            # Get dense candidates

            raw = self.vs.collection.query(query_embeddings=[q_emb.tolist()], n_results=DENSE_RETRIEVAL_CANDIDATES, include=["documents", "metadatas", "embeddings"])

            if not raw["documents"][0]:

                st.warning("No documents found in the vector store.")

                return []

            dense_cands = [{"id": raw["ids"][0][i], "content": raw["documents"][0][i], "metadata": raw["metadatas"][0][i], "emb": raw["embeddings"][0][i]} for i in range(len(raw["documents"][0]))]



            # Hybrid fusion: normalize and combine scores

            cands = self._hybrid_fusion(q_emb, sparse_cands, dense_cands)

        else:

            # Dense-only retrieval

            raw = self.vs.collection.query(query_embeddings=[q_emb.tolist()], n_results=DENSE_RETRIEVAL_CANDIDATES, include=["documents", "metadatas", "embeddings"])

            if not raw["documents"][0]:

                st.warning("No documents found in the vector store.")

                return []

            cands = [{"id": raw["ids"][0][i], "content": raw["documents"][0][i], "metadata": raw["metadatas"][0][i], "emb": raw["embeddings"][0][i]} for i in range(len(raw["documents"][0]))]



            # Add similarity scores

            sims = cosine_similarity([q_emb], [c["emb"] for c in cands])[0]

            for i, s in enumerate(sims): 

                cands[i]["similarity_score"] = float(s)



        if not cands:

            return []



        # Apply cross-encoder reranking if available

        if self.cross is not None:

            cands = self._rerank_with_cross_encoder(q, cands, top_k * CROSS_ENCODER_CANDIDATE_MULTIPLIER)



        # Return top_k results

        res = sorted(cands, key=lambda x: x.get("similarity_score", 0.0), reverse=True)[:top_k]

        return res



    def _hybrid_fusion(self, q_emb, sparse_cands, dense_cands):

        """Fuse sparse and dense retrieval results using reciprocal rank fusion"""

        # Create lookup for all candidates

        all_cands = {}



        # Normalize sparse scores

        sparse_scores = [c.get("sparse_score", 0.0) for c in sparse_cands]

        max_sparse = max(sparse_scores) if sparse_scores else 1.0



        for rank, c in enumerate(sparse_cands):

            cid = c.get("id", sha256_text_shared(c.get("content", "")))

            normalized_sparse = c.get("sparse_score", 0.0) / max(max_sparse, 1e-10)

            if cid not in all_cands:

                all_cands[cid] = c.copy()

                all_cands[cid]["sparse_rank"] = rank

                all_cands[cid]["sparse_score_norm"] = normalized_sparse

            else:

                all_cands[cid]["sparse_rank"] = rank

                all_cands[cid]["sparse_score_norm"] = normalized_sparse



        # Add dense scores

        dense_sims = cosine_similarity([q_emb], [c["emb"] for c in dense_cands])[0]

        for rank, (c, sim) in enumerate(zip(dense_cands, dense_sims)):

            cid = c.get("id", sha256_text_shared(c.get("content", "")))

            if cid not in all_cands:

                all_cands[cid] = c.copy()

                all_cands[cid]["dense_rank"] = rank

                all_cands[cid]["dense_score"] = float(sim)

            else:

                all_cands[cid]["dense_rank"] = rank

                all_cands[cid]["dense_score"] = float(sim)



        # Reciprocal rank fusion: score = sum(1 / (rank + k))

        for cid, c in all_cands.items():

            sparse_rank = c.get("sparse_rank", DEFAULT_MISSING_RANK)

            dense_rank = c.get("dense_rank", DEFAULT_MISSING_RANK)

            rrf_score = (1.0 / (sparse_rank + RRF_RANK_CONSTANT)) + (1.0 / (dense_rank + RRF_RANK_CONSTANT))



            # Also maintain individual scores for transparency

            sparse_score_norm = c.get("sparse_score_norm", 0.0)

            dense_score = c.get("dense_score", 0.0)



            # Final hybrid score with weighted combination

            c["similarity_score"] = float(

                HYBRID_FUSION_RRF_WEIGHT * rrf_score + 

                HYBRID_FUSION_DENSE_WEIGHT * dense_score + 

                HYBRID_FUSION_SPARSE_WEIGHT * sparse_score_norm

            )



        return list(all_cands.values())



    def _rerank_with_cross_encoder(self, query, candidates, top_n):

        """Rerank candidates using cross-encoder"""

        try:

            pairs = [[query, c["content"]] for c in candidates]

            scores = self.cross.predict(pairs)



            for i, score in enumerate(scores):

                # Keep original similarity_score, add cross_encoder_score

                candidates[i]["cross_encoder_score"] = float(score)

                # Update similarity_score to be weighted combination

                original_score = candidates[i].get("similarity_score", 0.0)

                candidates[i]["similarity_score"] = (

                    ORIGINAL_SCORE_WEIGHT * original_score + 

                    CROSS_ENCODER_WEIGHT * float(score)

                )



            return sorted(candidates, key=lambda x: x.get("similarity_score", 0.0), reverse=True)[:top_n]

        except Exception as e:

            st.warning(f"Cross-encoder reranking failed: {e}")

            return candidates



class GroqLLMRAG:

    def __init__(self, model="llama-3.1-8b-instant", key=None):

        self.llm = ChatGroq(groq_api_key=key, model_name=model, temperature=0.1)



    def generate_response(self, q, ctx, prompt_instr=""):

        # Note: Citation logic is handled separately by dedicated citation functions

        p = f"""{prompt_instr}
TASK: Answer the user's question using ONLY the provided context from tender documents.

CONTEXT:
{ctx}

QUESTION: {q}

INSTRUCTIONS:
1. Answer must be based EXCLUSIVELY on the context above - do not add external knowledge
2. If the context doesn't contain complete information, state: "The provided context contains partial information:" and explain what is available
3. If the context contains no relevant information, respond: "The provided context does not contain sufficient information to answer this question."
4. For tender-related questions, distinguish between mandatory requirements (must/shall) and optional items (should/may)
5. Use clear, professional language with short paragraphs
6. Each statement must be directly traceable to the context provided
7. Cite specific details from context when available (dates, numbers, requirements, page references)
8. If the question has multiple parts, address each part separately
9. Output format: Plain text response with NO markdown or special formatting unless present in original context

ANSWER:"""

        return self.llm.invoke([HumanMessage(content=p)]).content



    def expand_query(self, q):

        """Generate query variations to improve retrieval"""

        try:

            prompt = f"""Generate {MAX_QUERY_VARIATIONS} alternative phrasings of the following question for better document retrieval.

ORIGINAL QUESTION: "{q}"

REQUIREMENTS:
1. Each variation must preserve the core intent but use different vocabulary and phrasing
2. Vary sentence structure (e.g., interrogative, declarative, keyword-based forms)
3. Ensure semantic diversity - each should potentially match different document sections
4. Focus on different aspects or perspectives of the same question
5. Avoid overly generic or broad alternatives
6. Keep variations concise and focused

OUTPUT FORMAT:
Respond with ONLY the alternative questions, one per line, without numbering, bullets, or explanations.
Do NOT include the original question in your response.

VARIATIONS:"""



            response = self.llm.invoke([HumanMessage(content=prompt)]).content

            # Filter out lines that look like numbered/bulleted lists or are too short

            variations = []
            for line in response.strip().split('\n'):
                stripped = line.strip()
                if stripped and len(stripped) > 3 and not re.match(r'^[\d\w][\.\)\:]', stripped):
                    variations.append(stripped)



            # Return original query plus variations

            return [q] + variations[:MAX_QUERY_VARIATIONS]

        except Exception as e:

            st.warning(f"Query expansion failed: {e}")

            return [q]  # Fallback to original query



    def validate_answer(self, q, answer, context):

        """Validate that the answer is grounded in the context"""

        try:

            prompt = f"""Validate if the answer is accurately grounded in the provided context.

QUESTION: {q}

PROVIDED ANSWER: {answer}

SOURCE CONTEXT: {context[:ANSWER_VALIDATION_CONTEXT_CHARS]}

VALIDATION CRITERIA:
1. Check if every factual claim in the answer is supported by the context
2. Identify any statements made without clear context support
3. Flag if the answer adds interpretation or inference beyond what's stated in context
4. Verify that the answer doesn't contradict any information in the context
5. Assess if critical context information is missing from the answer

RESPONSE FORMAT:
Start with 'YES' if the answer is well-grounded, or 'NO' if it has significant issues.
Then provide a brief explanation covering:
- What parts are well-supported
- Any unsupported claims or issues found
- Overall assessment

VALIDATION:"""



            response = self.llm.invoke([HumanMessage(content=prompt)]).content

            is_valid = response.strip().upper().startswith('YES')

            return is_valid, response

        except Exception:

            return True, "Validation skipped due to error"





def assemble_context_rag(chunks, max_ctx=4096):

    parts = [f"Source: {c['metadata'].get('source_file')} (Page: {c['metadata'].get('page', 'N/A')}) [ID: {c['id'][:8]}]\n{c['content']}" for c in chunks]

    full_text = "\n\n".join(parts)

    # Enforce max_ctx by truncating if necessary

    if len(full_text) > max_ctx * CHARS_PER_TOKEN_ESTIMATE:

        full_text = full_text[:max_ctx * CHARS_PER_TOKEN_ESTIMATE]

    return full_text, chunks



def multi_query_retrieval(retriever, queries, top_k=5):

    """Retrieve documents using multiple query variations and fuse results"""

    all_results = {}



    for query_idx, query in enumerate(queries):

        results = retriever.retrieve(query, top_k=top_k * MULTI_QUERY_RETRIEVAL_MULTIPLIER)



        for rank, result in enumerate(results):

            doc_id = result.get("id", sha256_text_shared(result.get("content", "")))



            if doc_id not in all_results:

                all_results[doc_id] = result.copy()

                all_results[doc_id]["query_ranks"] = []

                all_results[doc_id]["query_scores"] = []



            all_results[doc_id]["query_ranks"].append(rank)

            all_results[doc_id]["query_scores"].append(result.get("similarity_score", 0.0))



    # Fusion: Use reciprocal rank fusion across queries

    for doc_id, result in all_results.items():

        ranks = result["query_ranks"]

        scores = result["query_scores"]



        # RRF score

        rrf = sum(1.0 / (r + RRF_RANK_CONSTANT) for r in ranks)

        # Average score across queries

        avg_score = sum(scores) / len(scores) if scores else 0.0



        # Combined score with weighted fusion

        result["similarity_score"] = MULTI_QUERY_RRF_WEIGHT * rrf + MULTI_QUERY_AVG_WEIGHT * avg_score

        result["num_queries_matched"] = len(ranks)



    # Sort and return top_k

    sorted_results = sorted(all_results.values(), 

                          key=lambda x: x.get("similarity_score", 0.0), 

                          reverse=True)[:top_k]



    return sorted_results





def extract_sentences_from_chunks(chunks):

    """Extract sentences from chunks with their positions and metadata"""

    all_sentences = []

    for chunk_idx, chunk in enumerate(chunks):

        content = chunk.get("content", "")

        sentences = simple_sent_tokenize_shared(content)



        char_pos = 0

        for sent in sentences:

            sent_start = content.find(sent, char_pos)

            if sent_start == -1:

                sent_start = char_pos



            all_sentences.append({

                "text": sent,

                "chunk_idx": chunk_idx,

                "chunk_id": chunk.get("id", ""),

                "source_file": chunk.get("metadata", {}).get("source_file", ""),

                "page": chunk.get("metadata", {}).get("page", "N/A"),

                "char_start": sent_start,

                "char_end": sent_start + len(sent)

            })

            char_pos = sent_start + len(sent)



    return all_sentences



def map_sentences_to_sources_rag(answer, used, em):

    """Map answer sentences to specific source sentences with precise citations"""

    answer_sents = simple_sent_tokenize_shared(answer)

    if not answer_sents or not used: 

        return []



    try:

        # Extract all sentences from source chunks

        source_sents = extract_sentences_from_chunks(used)

        if not source_sents:

            return []



        # Generate embeddings

        answer_embs = em.generate_embeddings(answer_sents)

        source_texts = [s["text"] for s in source_sents]

        source_embs = em.generate_embeddings(source_texts)



        # Compute similarities

        sims = cosine_similarity(answer_embs, source_embs)



        # Map each answer sentence to best matching source sentence

        citations = []

        for i, ans_sent in enumerate(answer_sents):

            best_idx = int(np.argmax(sims[i]))

            best_score = float(sims[i][best_idx])

            best_source = source_sents[best_idx]



            citations.append({

                "answer_sentence": ans_sent,

                "source_sentence": best_source["text"],

                "source_file": best_source["source_file"],

                "page": best_source["page"],

                "chunk_id": best_source["chunk_id"],

                "similarity_score": best_score,

                "char_range": f"{best_source['char_start']}-{best_source['char_end']}"

            })



        return citations

    except Exception as e:

        st.warning(f"Failed to map sentences to sources: {e}")

        return []



def extract_clauses_from_chunks(chunks):

    """Extract clause-level segments from chunks for finer-grained citations"""

    clauses = []



    # Use module-level clause patterns

    combined_pattern = '|'.join(f'({p})' for p in CLAUSE_PATTERNS)



    for chunk_idx, chunk in enumerate(chunks):

        content = chunk.get("content", "")



        # Split by clause markers

        parts = re.split(combined_pattern, content)



        # Reconstruct clauses with their markers

        current_clause = ""

        char_pos = 0



        for part in parts:

            if part and part.strip():

                current_clause += part



                # If clause is substantial, add it

                if len(current_clause.strip()) > MIN_CLAUSE_LENGTH:

                    clause_start = content.find(current_clause.strip(), char_pos)

                    if clause_start == -1:

                        clause_start = char_pos



                    clauses.append({

                        "text": current_clause.strip(),

                        "chunk_idx": chunk_idx,

                        "chunk_id": chunk.get("id", ""),

                        "source_file": chunk.get("metadata", {}).get("source_file", ""),

                        "page": chunk.get("metadata", {}).get("page", "N/A"),

                        "char_start": clause_start,

                        "char_end": clause_start + len(current_clause.strip())

                    })



                    char_pos = clause_start + len(current_clause.strip())

                    current_clause = ""



        # Add any remaining clause

        if current_clause.strip() and len(current_clause.strip()) > MIN_CLAUSE_LENGTH:

            clauses.append({

                "text": current_clause.strip(),

                "chunk_idx": chunk_idx,

                "chunk_id": chunk.get("id", ""),

                "source_file": chunk.get("metadata", {}).get("source_file", ""),

                "page": chunk.get("metadata", {}).get("page", "N/A"),

                "char_start": char_pos,

                "char_end": char_pos + len(current_clause.strip())

            })



    return clauses



def map_answer_to_clauses_rag(answer, used, em):

    """Map answer to clause-level citations"""

    answer_sents = simple_sent_tokenize_shared(answer)

    if not answer_sents or not used:

        return []



    try:

        # Extract clauses from source chunks

        clauses = extract_clauses_from_chunks(used)

        if not clauses:

            return []



        # Generate embeddings

        answer_embs = em.generate_embeddings(answer_sents)

        clause_texts = [c["text"] for c in clauses]

        clause_embs = em.generate_embeddings(clause_texts)



        # Compute similarities

        sims = cosine_similarity(answer_embs, clause_embs)



        # Map each answer sentence to best matching clause

        clause_citations = []

        for i, ans_sent in enumerate(answer_sents):

            best_idx = int(np.argmax(sims[i]))

            best_score = float(sims[i][best_idx])

            best_clause = clauses[best_idx]



            # Truncate clause text for display if too long

            clause_text = best_clause["text"]

            if len(clause_text) > MAX_CLAUSE_DISPLAY_LENGTH:

                clause_text = clause_text[:MAX_CLAUSE_DISPLAY_LENGTH] + "..."



            clause_citations.append({

                "answer_sentence": ans_sent,

                "source_clause": clause_text,

                "source_file": best_clause["source_file"],

                "page": best_clause["page"],

                "chunk_id": best_clause["chunk_id"],

                "similarity_score": best_score,

                "char_range": f"{best_clause['char_start']}-{best_clause['char_end']}"

            })



        return clause_citations

    except Exception as e:

        st.warning(f"Failed to map answer to clauses: {e}")

        return []



# --------------------------

# Summarization Components

# --------------------------

GLOBAL_SUM_CONCURRENCY = asyncio.Semaphore(2)



def clean_json_string(txt):
    """Clean and fix common JSON formatting issues."""
    
    # Remove any leading/trailing whitespace
    txt = txt.strip()
    
    # Extract JSON from markdown code block if present
    if "```json" in txt:
        try:
            parts = txt.split("```json", 1)
            if len(parts) > 1:
                json_parts = parts[1].split("```", 1)
                if len(json_parts) > 0:
                    txt = json_parts[0].strip()
        except (IndexError, ValueError, AttributeError):
            pass
    
    # Remove any "```" markers that might remain
    txt = txt.replace("```", "").strip()
    
    # Try to fix common JSON issues
    # 1. Replace single quotes with double quotes for keys and values
    # Using a more inclusive pattern to match keys with hyphens, spaces, etc.
    # Pattern: 'key': 'value' -> "key": "value"
    txt = re.sub(r"'([^']+)'(\s*:\s*)'([^']*)'", r'"\1"\2"\3"', txt)
    txt = re.sub(r"'([^']+)'(\s*:\s*)", r'"\1"\2', txt)  # 'key': -> "key":
    
    # 2. Remove trailing commas before } or ]
    txt = re.sub(r',(\s*[}\]])', r'\1', txt)
    
    return txt


async def call_gemini_json_sum_async(client, sys, user, model, rpm):

    await asyncio.sleep(DEFAULT_RATE_LIMIT_SLEEP / max(1, rpm))  # Proper rate limiting

    async with GLOBAL_SUM_CONCURRENCY:

        try:

            from google.genai import types

            resp = await client.aio.models.generate_content(model=model, contents=user, config=types.GenerateContentConfig(system_instruction=sys, response_mime_type="application/json"))

            txt = resp.text or "{}"

            # Clean the JSON string
            txt = clean_json_string(txt)

            # First attempt: parse as-is
            try:
                parsed = json.loads(txt)
                return parsed
            except json.JSONDecodeError:
                # Second attempt: try replacing all single quotes with double quotes
                # WARNING: This is aggressive and will break valid strings with apostrophes
                # (e.g., "it's" becomes "it"s"). This is a known limitation.
                txt_fixed = txt.replace("'", '"')
                try:
                    parsed = json.loads(txt_fixed)
                    return parsed
                except json.JSONDecodeError as e:
                    # If it still fails, return the error
                    return {"error": f"Invalid JSON response: {str(e)}"}

        except Exception as e:

            return {"error": f"API call failed: {str(e)}"}



def extract_pages_sum(path):

    doc = None

    try:

        doc = fitz.open(path)

        p, f = [], []

        # pdfplumber context manager handles its own cleanup

        with pdfplumber.open(path) as pl_doc:

            for i, page in enumerate(doc):

                t = page.get_text()

                if i < len(pl_doc.pages):

                    tabs = pl_doc.pages[i].extract_tables()

                    for tab in tabs:

                        if tab: t += f"\n\n[TABLE]\n{pd.DataFrame(tab).to_markdown()}\n"

                p.append(t); f.append(len(page.get_images()) > 0)

        return p, f

    finally:

        if doc:

            doc.close()



def semantic_chunk_pages_sum(pages, flags, max_tok=8000, overlap=5):

    flat = []

    for i, p in enumerate(pages):

        if _NLTK_READY:

            try:

                sents = nltk.sent_tokenize(p)

            except Exception:

                # Fallback to simple sentence tokenization if NLTK fails

                sents = simple_sent_tokenize_shared(p)

        else:

            sents = simple_sent_tokenize_shared(p)

        for s in sents:

            flat.append((i+1, s, flags[i]))

    chunks, bs, bp, bf = [], [], [], []

    for p, s, f in flat:

        if sum(len(x)//4 for x in bs + [s]) > max_tok and bs:

            chunks.append({"text": " ".join(bs).strip(), "start_page": min(bp), "has_visual": any(bf), "id": str(uuid.uuid4())[:8]})

            bs, bp, bf = bs[-overlap:] if overlap > 0 else [], bp[-overlap:] if overlap > 0 else [], bf[-overlap:] if overlap > 0 else []

        bs.append(s); bp.append(p); bf.append(f)

    if bs: chunks.append({"text": " ".join(bs).strip(), "start_page": min(bp), "has_visual": any(bf), "id": str(uuid.uuid4())[:8]})

    return [c for c in chunks if c["text"].strip()]



def get_sum_prompts(mode: str):

    if mode == "Compliance Matrix":

        map_system = """You are a compliance requirements extractor analyzing tender documents. Extract specific, actionable requirements."""
        
        map_instruction = """Extract ALL compliance requirements, obligations, and conditions from the document chunk.

EXTRACTION RULES:
1. Identify explicit requirements (must/shall/required statements)
2. Include implicit requirements (should/recommended items)
3. Capture specific criteria (thresholds, dates, qualifications)
4. Note any conditional requirements (if-then clauses)
5. Reference the page number where each requirement appears

OUTPUT FORMAT (strict JSON array):
[
  {
    "item": "Brief requirement title",
    "detail": "Complete requirement description with specific criteria",
    "category": "One of: Financial, Technical, Legal, Insurance, Operational, Other",
    "mandatory": true/false,
    "page": 1
  }
]

If no requirements found in this chunk, return: []"""
        
        reduce_system = """You are consolidating compliance requirements from multiple document chunks into a unified matrix."""
        
        reduce_instruction = """You will receive mapped findings from document analysis in the 'D:' parameter below.
USE THIS DATA to create your comprehensive compliance matrix. The D: parameter contains an array of findings from the map phase.
If D: contains empty arrays or no requirements, still generate the report structure with empty matrix array and appropriate summary statistics.

Merge all extracted requirements into a comprehensive compliance matrix.

CONSOLIDATION RULES:
1. Remove exact duplicates (same requirement appearing multiple times)
2. Merge similar requirements with all page references
3. Group by category for better organization
4. Preserve all unique mandatory vs optional distinctions
5. Maintain complete detail for each requirement

OUTPUT FORMAT (strict JSON):
{
  "matrix": [
    {
      "item": "Requirement title",
      "detail": "Full description",
      "category": "...",
      "mandatory": true/false,
      "pages": [1, 5, 12]
    }
  ],
  "total_requirements": 0,
  "summary": {
    "mandatory_count": 0,
    "optional_count": 0,
    "categories": {}
  }
}"""

        return (map_system, map_instruction, reduce_system, reduce_instruction)

    elif mode == "Risk Assessment":

        map_system = """You are a risk analyst identifying risks, liabilities, and concerns in tender documents."""
        
        map_instruction = """Extract ALL identified or potential risks from the document chunk.

RISK IDENTIFICATION:
1. Look for explicit risk mentions (penalties, liabilities, constraints)
2. Identify implicit risks (tight deadlines, complex requirements, ambiguous terms)
3. Note financial risks (payment terms, penalties, insurance requirements)
4. Flag operational risks (resource constraints, dependencies)
5. Include legal/compliance risks (regulatory requirements, liability clauses)

OUTPUT FORMAT (strict JSON array):
[
  {
    "clause": "Brief description of the risky clause or requirement",
    "risk_level": "One of: Critical, High, Medium, Low",
    "risk_type": "One of: Financial, Operational, Legal, Reputational, Technical",
    "impact": "Description of potential impact",
    "page": 1
  }
]

If no risks identified, return: []"""
        
        reduce_system = """You are consolidating risk assessments from multiple document chunks."""
        
        reduce_instruction = """You will receive mapped findings from document analysis in the 'D:' parameter below.
USE THIS DATA to create your comprehensive risk assessment. The D: parameter contains an array of findings from the map phase.
If D: contains empty arrays or no risks, still generate the report structure with empty risks array and appropriate summary statistics.

Merge all identified risks into a comprehensive risk assessment.

CONSOLIDATION RULES:
1. Remove duplicate risk entries
2. Merge related risks with all page references
3. Prioritize by risk level (Critical > High > Medium > Low)
4. Group by risk type for strategic overview
5. Provide risk distribution statistics

OUTPUT FORMAT (strict JSON):
{
  "risks": [
    {
      "clause": "Risk description",
      "risk_level": "...",
      "risk_type": "...",
      "impact": "...",
      "pages": [2, 8, 15]
    }
  ],
  "risk_summary": {
    "total_risks": 0,
    "critical_count": 0,
    "high_count": 0,
    "medium_count": 0,
    "low_count": 0,
    "by_type": {}
  }
}"""

        return (map_system, map_instruction, reduce_system, reduce_instruction)

    elif mode == "Entity Dashboard":

        map_system = """You are extracting key entities and metadata from tender documents."""
        
        map_instruction = """Extract ALL named entities and important metadata from the document chunk.

ENTITY CATEGORIES:
1. Organizations: Companies, agencies, departments
2. People: Key contacts, decision-makers
3. Locations: Addresses, sites, regions
4. Dates: Deadlines, milestones, timeframes
5. Financial: Budgets, amounts, payment terms
6. Technical: Systems, standards, specifications

OUTPUT FORMAT (strict JSON array):
[
  {
    "category": "One of: Organizations, People, Locations, Dates, Financial, Technical",
    "entity": "The specific entity name or value",
    "context": "Brief context about relevance",
    "page": 1
  }
]

If no entities found, return: []"""
        
        reduce_system = """You are compiling entities into an organized dashboard."""
        
        reduce_instruction = """You will receive mapped findings from document analysis in the 'D:' parameter below.
USE THIS DATA to create your structured entity dashboard. The D: parameter contains an array of findings from the map phase.
If D: contains empty arrays or no entities, still generate the report structure with empty dashboard categories and appropriate summary statistics.

Consolidate all entities into a structured dashboard.

COMPILATION RULES:
1. Group entities by category
2. Remove duplicate entities (keep all page references)
3. Organize chronologically for dates
4. Prioritize high-value/critical entities
5. Create summary statistics

OUTPUT FORMAT (strict JSON):
{
  "dashboard": {
    "Organizations": [],
    "People": [],
    "Locations": [],
    "Dates": [],
    "Financial": [],
    "Technical": []
  },
  "entity_count": {
    "Organizations": 0,
    "People": 0,
    "Locations": 0,
    "Dates": 0,
    "Financial": 0,
    "Technical": 0
  }
}"""

        return (map_system, map_instruction, reduce_system, reduce_instruction)

    elif mode == "Ambiguity Scrutiny":

        map_system = """You are analyzing tender documents to identify ambiguous, vague, or unclear language that could lead to misinterpretation."""
        
        map_instruction = """Extract ALL instances of ambiguity, vagueness, or unclear language from the document chunk.

AMBIGUITY IDENTIFICATION:
1. Look for vague terms (e.g., "reasonable", "appropriate", "timely", "sufficient" without clear definitions)
2. Identify unclear requirements or conditions that lack specific criteria
3. Flag conflicting or contradictory statements
4. Note undefined terms or acronyms used without explanation
5. Highlight incomplete information or missing details
6. Identify subjective language without objective measures

OUTPUT FORMAT (strict JSON array):
[
  {
    "ambiguous_text": "The exact text that is ambiguous or unclear",
    "ambiguity_type": "One of: Vague Terms, Unclear Requirements, Conflicting Statements, Undefined Terms, Missing Information, Subjective Language",
    "issue": "Clear explanation of what makes this ambiguous",
    "potential_impact": "How this ambiguity could affect compliance or interpretation",
    "severity": "One of: High, Medium, Low",
    "page": 1
  }
]

If no ambiguities found in this chunk, return: []"""
        
        reduce_system = """You are consolidating ambiguity findings from multiple document chunks into a comprehensive report."""
        
        reduce_instruction = """You will receive mapped findings from document analysis in the 'D:' parameter below.
USE THIS DATA to create your comprehensive scrutiny report. The D: parameter contains an array of findings from the map phase.
If D: contains empty arrays or no ambiguity findings, still generate the report structure with empty ambiguities array and appropriate summary statistics.

Merge all identified ambiguities into a comprehensive scrutiny report.

CONSOLIDATION RULES:
1. Remove duplicate ambiguities (same issue appearing multiple times)
2. Merge similar ambiguities with all page references
3. Prioritize by severity (High > Medium > Low)
4. Group by ambiguity type for systematic review
5. Provide actionable recommendations for each ambiguity

OUTPUT FORMAT (strict JSON):
{
  "ambiguities": [
    {
      "ambiguous_text": "The ambiguous text",
      "ambiguity_type": "...",
      "issue": "Explanation of the ambiguity",
      "potential_impact": "Impact description",
      "severity": "...",
      "pages": [1, 5, 12],
      "recommendation": "Suggested clarification or action"
    }
  ],
  "summary": {
    "total_ambiguities": 0,
    "high_severity": 0,
    "medium_severity": 0,
    "low_severity": 0,
    "by_type": {}
  },
  "overall_assessment": "Brief assessment of the document's clarity and potential risks from ambiguities"
}"""

        return (map_system, map_instruction, reduce_system, reduce_instruction)

    else:  # General Summary

        map_system = """You are extracting key information and findings from tender documents."""
        
        map_instruction = """Extract important findings, details, and insights from the document chunk.

EXTRACTION FOCUS:
1. Main topics and themes covered
2. Critical requirements or obligations
3. Important facts, figures, and data points
4. Notable conditions or constraints
5. Any unique or unusual provisions

OUTPUT FORMAT (strict JSON array):
[
  {
    "finding": "Clear, concise description of the finding",
    "type": "One of: Requirement, Condition, Fact, Constraint, Opportunity",
    "importance": "One of: High, Medium, Low",
    "page": 1
  }
]

If chunk contains no significant information, return: []"""
        
        reduce_system = """You are synthesizing findings into a comprehensive summary."""
        
        reduce_instruction = """You will receive mapped findings from document analysis in the 'D:' parameter below.
USE THIS DATA to create your comprehensive summary. The D: parameter contains an array of findings from the map phase.
If D: contains empty arrays or no findings, still generate the report structure with empty key_findings array and appropriate summary statistics.

Create a well-structured summary from all extracted findings.

SYNTHESIS RULES:
1. Organize findings by type and importance
2. Create a coherent narrative summary
3. Highlight high-importance items
4. Include all relevant page citations
5. Provide statistics on findings

OUTPUT FORMAT (strict JSON):
{
  "summary": "Comprehensive 2-3 paragraph summary of key findings",
  "key_findings": [
    {
      "finding": "...",
      "type": "...",
      "importance": "...",
      "pages": [1, 3, 7]
    }
  ],
  "citations": [1, 2, 3, 5, 7],
  "finding_stats": {
    "total_findings": 0,
    "high_importance": 0,
    "by_type": {}
  }
}"""

        return (map_system, map_instruction, reduce_system, reduce_instruction)



def compute_confidence_score_sum(mapped, reduced, q):

    c = {"snippet_coverage": min(1.0, len(mapped)/5.0), "result_coherence": 0.8 if isinstance(reduced, dict) and "error" not in reduced else 0.1, "information_density": 0.5, "citation_confidence": 0.7}

    c["overall_confidence"] = 0.25*c["snippet_coverage"] + 0.25*c["result_coherence"] + 0.25*c["information_density"] + 0.25*c["citation_confidence"]

    return c



def render_page_level_citations_sum(chunks, pages):
    """Display page-level citation tracking with text snippets from chunks"""
    
    if not chunks or not pages:
        return
    
    st.markdown("---")
    with st.expander("📋 Citation Tracking: Page-Level Extraction", expanded=True):
        st.markdown("**Source pages and text excerpts that informed this analysis:**")
        
        # Group chunks by page
        page_chunks = {}
        for chunk in chunks:
            page = chunk.get('start_page', 0)
            if page in pages:
                if page not in page_chunks:
                    page_chunks[page] = []
                page_chunks[page].append(chunk)
        
        # Display citations for each page
        for page_num in sorted(page_chunks.keys()):
            st.markdown(f"#### 📄 Page {page_num}")
            
            for idx, chunk in enumerate(page_chunks[page_num]):
                text = chunk.get('text', '')
                chunk_id = chunk.get('id', 'N/A')
                
                # Truncate text for display
                display_text = text
                if len(display_text) > MAX_CLAUSE_DISPLAY_LENGTH:
                    display_text = display_text[:MAX_CLAUSE_DISPLAY_LENGTH] + "..."
                
                st.markdown(f"**Citation {idx + 1}** (Chunk ID: `{chunk_id}`)")
                st.text(display_text)
                
                # Show full text in a nested expander if truncated
                if len(text) > MAX_CLAUSE_DISPLAY_LENGTH:
                    with st.expander(f"Show full text for citation {idx + 1}"):
                        st.text(text)
            
            st.markdown("---")


def render_citation_preview_sum(doc, cites):

    if not cites or not doc: 

        return

    st.markdown("### 📄 Context Preview")

    # Show all pages, not just first 5
    tabs = st.tabs([f"Page {c['page']}" for c in cites])

    for i, c in enumerate(cites):

        with tabs[i]:

            try:

                page = doc.load_page(c['page'] - 1)

                pix = page.get_pixmap(matrix=fitz.Matrix(ZOOM_LEVEL, ZOOM_LEVEL))

                st.image(pix.tobytes("png"), caption=f"Page {c['page']}", use_container_width=True)

            except Exception as e:

                st.error(f"Failed to load page {c['page']}: {e}")



def _format_pages_column(df):
    """Helper function to convert pages list to string for better display"""
    if "pages" in df.columns:
        df["pages"] = df["pages"].apply(lambda x: ", ".join(map(str, x)) if isinstance(x, list) else str(x))
    return df


def sanitize_filename(text, max_length=30):
    """Sanitize text for use in filenames by removing invalid characters"""
    # Remove or replace invalid filename characters
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        text = text.replace(char, '_')
    # Limit length and strip whitespace
    return text[:max_length].strip().replace(' ', '_')


def convert_result_to_dataframe(result, objective):
    """Convert summarization result JSON to pandas DataFrame based on objective type"""
    
    if not isinstance(result, dict):
        return None
    
    try:
        df = None
        
        if objective == "Compliance Matrix" and "matrix" in result:
            df = pd.DataFrame(result["matrix"])
        
        elif objective == "Risk Assessment" and "risks" in result:
            df = pd.DataFrame(result["risks"])
        
        elif objective == "Entity Dashboard" and "dashboard" in result:
            # Flatten the nested dashboard structure into a DataFrame
            entities = []
            for category, items in result["dashboard"].items():
                if isinstance(items, list):
                    for item in items:
                        if isinstance(item, dict):
                            # Shallow copy to avoid mutating the original result dictionary
                            item_copy = item.copy()
                            item_copy["category"] = category
                            entities.append(item_copy)
            
            # Create DataFrame from entities
            if entities:
                df = pd.DataFrame(entities)
            else:
                # Create empty DataFrame with expected columns for consistent output
                df = pd.DataFrame(columns=["category", "entity", "context", "page"])
        
        elif objective == "Ambiguity Scrutiny" and "ambiguities" in result:
            df = pd.DataFrame(result["ambiguities"])
        
        elif objective == "General Summary":
            # For general summary, try to extract any list data
            for key in result.keys():
                if isinstance(result[key], list) and result[key]:
                    df = pd.DataFrame(result[key])
                    break
        
        # Format pages column if DataFrame was created
        if df is not None:
            df = _format_pages_column(df)
        
        return df
    except Exception as e:
        st.warning(f"Could not convert result to table: {e}")
        return None


def export_to_excel(df, query, objective):
    """Export DataFrame to Excel file and return bytes"""
    
    if df is None or df.empty:
        return None
    
    try:
        from io import BytesIO
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Results')
            
            # Get the worksheet
            worksheet = writer.sheets['Results']
            
            # Auto-adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column = [cell for cell in column]
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except (TypeError, AttributeError):
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column[0].column_letter].width = adjusted_width
        
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        st.error(f"Error creating Excel file: {e}")
        return None


def export_to_word(df, query, objective):
    """Export DataFrame to Word document and return bytes"""
    
    if df is None or df.empty:
        return None
    
    try:
        from io import BytesIO
        output = BytesIO()
        
        doc = Document()
        doc.add_heading(f'{objective} - Results', 0)
        doc.add_paragraph(f'Query: {query}')
        doc.add_paragraph('')
        
        # Add table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            header_cells[i].text = str(column_name)
        
        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if value is not None else ''
        
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        st.error(f"Error creating Word document: {e}")
        return None


# --------------------------

# Main Program logic

# --------------------------

st.sidebar.title("App Navigation")

choice = st.sidebar.radio("Choose Program", ["Simple QA (RAG)", "Context Understanding (Summarization)"])



if choice == "Simple QA (RAG)":

    st.title("Tender Analyser-Simple QA Mode")

    with st.sidebar:

        st.header("RAG Settings")

        ga_key = st.text_input("Groq API Key", type="password", key="rgk")

        emb_m = st.text_input("Embedding model", "all-MiniLM-L6-v2", key="rem")

        crs_m = st.text_input("Cross-Encoder", "cross-encoder/ms-marco-MiniLM-L-6-v2", key="rcm")

        u_crs = st.checkbox("Use Cross-Encoder", True, key="rucx")

        u_spa = st.checkbox("Use Sparse BM25", True, key="rusp")



        # Chunking method selection

        chunk_method = st.selectbox(

            "Chunking Method",

            ["Recursive Character Splitter", "Hybrid (Token+Semantic)", "Fixed Context Window"],

            key="chunk_method"

        )



        c_size = st.number_input("Chunk size", 800, key="rcsi")

        c_over = st.number_input("Overlap", 128, key="rcov")

        max_tk = st.number_input("Context Tokens", 4096, key="rctx")

        top_k = st.number_input("Top K UI", 3, key="rtk")

        supp_th = st.slider("Support threshold", 0.1, 1.0, 0.45, key="rsth")



    files = st.file_uploader("Upload PDFs", type="pdf", accept_multiple_files=True, key="rfl")

    if files:

        # Validate file sizes

        for f in files:

            file_size_mb = len(f.getvalue()) / (1024 * 1024)

            if file_size_mb > MAX_FILE_SIZE_MB:

                st.error(f"File {f.name} exceeds maximum size of {MAX_FILE_SIZE_MB}MB (size: {file_size_mb:.1f}MB)")

                st.stop()



        if st.button("Index / Re-index", key="rib"):

            with st.spinner("Indexing..."):

                em = EmbeddingManagerRAG(emb_m); vs = VectorStoreRAG()

                all_chunks = []

                for f in files:

                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as t:

                        try:

                            t.write(f.getbuffer())

                            t.flush()

                            h = sha256_file_shared(t.name)



                            # Load PDF using PyMuPDF

                            docs = load_pdf_with_pymupdf(t.name)



                            # Update metadata

                            for d in docs:

                                d.metadata.update({"file_hash": h, "source_file": f.name})



                            # Apply selected chunking method

                            chunks = apply_chunking_method(docs, chunk_method, c_size, c_over)

                            all_chunks.extend(chunks)

                        finally:

                            # Ensure temporary file is deleted

                            # Failures are acceptable here as the OS may have already cleaned up

                            try:

                                os.unlink(t.name)

                            except Exception:

                                pass  # File may already be deleted



                if not all_chunks:

                    st.error("No content could be extracted from the uploaded files.")

                    st.stop()



                texts = [c.page_content for c in all_chunks]

                embs = em.generate_embeddings(texts)

                vs.add_documents(all_chunks, embs)

                st.session_state["rvs"], st.session_state["sem"] = vs, em



                if u_spa and BM25_AVAILABLE:

                    bm = SparseBM25IndexRAG()

                    bm.build(all_chunks, embs)

                    st.session_state["rbm"] = bm

                elif u_spa and not BM25_AVAILABLE:

                    st.warning("BM25 is not available. Sparse retrieval disabled. Install: pip install rank-bm25")



                st.success("Indexing Success!")



    if "rvs" in st.session_state:

        q = st.text_input("Ask a question:", key="rqi")



        # Add advanced options

        with st.sidebar:

            st.markdown("---")

            st.subheader("Advanced Options")

            use_query_expansion = st.checkbox("Enable Query Expansion", True, key="rqe")

            use_answer_validation = st.checkbox("Validate Answer", False, key="rav")

            show_clause_citations = st.checkbox("Show Clause-Level Citations", True, key="rcc")

            show_sentence_citations = st.checkbox("Show Sentence-Level Citations", True, key="rsc")



        if q:

            # Validate query length

            if len(q) > MAX_QUERY_LENGTH:

                st.error(f"Query exceeds maximum length of {MAX_QUERY_LENGTH} characters")

                st.stop()



            if not ga_key:

                st.error("Please provide a Groq API Key in the sidebar")

                st.stop()



            with st.spinner("Thinking..."):


                # Start total time measurement


                total_start_time = time.time()


                

                # Cache cross-encoder model in session state

                if u_crs and CROSS_ENCODER_AVAILABLE:

                    if "cross_encoder" not in st.session_state:

                        st.session_state["cross_encoder"] = CrossEncoder(crs_m)

                    cr = st.session_state["cross_encoder"]

                else:

                    if u_crs and not CROSS_ENCODER_AVAILABLE:

                        st.warning("Cross-encoder not available. Install: pip install sentence-transformers")

                    cr = None



                # Initialize LLM and retriever

                llm = GroqLLMRAG(key=ga_key)

                ret = RAGRetrieverRAG(st.session_state["rvs"], st.session_state["sem"], cr, sparse=st.session_state.get("rbm"), use_s=u_spa)




                # Start retrieval time measurement


                retrieval_start_time = time.time()


                

                # Query expansion if enabled

                if use_query_expansion:

                    with st.spinner("Expanding query..."):

                        queries = llm.expand_query(q)

                        if len(queries) > 1:

                            st.info(f"Using {len(queries)} query variations for improved retrieval")

                            with st.expander("Query Variations"):

                                for i, qvar in enumerate(queries):

                                    st.write(f"{i+1}. {qvar}")



                    # Multi-query retrieval

                    res = multi_query_retrieval(ret, queries, top_k=top_k)

                else:

                    # Standard retrieval

                    res = ret.retrieve(q, top_k)




                # End retrieval time measurement


                retrieval_time = time.time() - retrieval_start_time


                

                if not res:

                    st.warning("No relevant documents found for your query. Try rephrasing or check if documents are properly indexed.")

                    st.stop()



                # Display retrieved context info

                with st.expander("📚 Retrieved Context Chunks"):

                    for i, chunk in enumerate(res):

                        st.markdown(f"**Chunk {i+1}** (Score: {chunk.get('similarity_score', 0):.3f})")

                        st.markdown(f"*Source: {chunk['metadata'].get('source_file')} | Page: {chunk['metadata'].get('page', 'N/A')}*")

                        # Show preview with truncation if needed

                        content = chunk['content']

                        if len(content) > MAX_CHUNK_PREVIEW_LENGTH:

                            content = content[:MAX_CHUNK_PREVIEW_LENGTH] + "..."

                        st.text(content)

                        st.markdown("---")




                # Start generation time measurement


                generation_start_time = time.time()


                

                # Generate answer

                ctx, used = assemble_context_rag(res, max_ctx=max_tk)

                ans = llm.generate_response(q, ctx)




                # End generation time measurement


                generation_time = time.time() - generation_start_time


                


                # Calculate total time


                total_time = time.time() - total_start_time


                

                # Answer validation if enabled

                if use_answer_validation:

                    with st.spinner("Validating answer..."):

                        is_valid, validation_msg = llm.validate_answer(q, ans, ctx)

                        if not is_valid:

                            st.warning(f"⚠️ Answer validation: {validation_msg}")



                # Display answer

                st.subheader("Answer")

                st.write(strip_provenance(ans))




                # Display timing metrics


                st.markdown("---")


                st.subheader("⏱️ Performance Metrics")


                col1, col2, col3 = st.columns(3)


                with col1:


                    st.metric("Retrieval Time", f"{retrieval_time:.2f} sec")


                with col2:


                    st.metric("Generation Time", f"{generation_time:.2f} sec")


                with col3:


                    st.metric("Total Time", f"{total_time:.2f} sec")


                

                # Confidence score

                supp = compute_mean_support_score_shared(res)

                if supp < supp_th:

                    st.warning(f"Low confidence support: {supp:.2f}")

                else:

                    st.success(f"Strong support: {supp:.2f}")



                # Clause-level citations

                if show_clause_citations:

                    clause_map = map_answer_to_clauses_rag(strip_provenance(ans), used, st.session_state["sem"])

                    if clause_map:

                        with st.expander("📋 Clause-Level Citations", expanded=True):

                            for i, m in enumerate(clause_map):

                                st.markdown(f"**Answer Sentence {i+1}:** {m['answer_sentence']}")

                                st.markdown(f"**Source Clause:** {m['source_clause']}")

                                st.caption(f"Source: {m['source_file']} | Page: {m['page']} | Score: {m['similarity_score']:.3f} | Chars: {m['char_range']}")

                                st.markdown("---")



                # Sentence-level citations

                if show_sentence_citations:

                    sent_map = map_sentences_to_sources_rag(strip_provenance(ans), used, st.session_state["sem"])

                    if sent_map:

                        with st.expander("📝 Sentence-Level Citations"):

                            for i, m in enumerate(sent_map):

                                st.markdown(f"**Answer:** {m['answer_sentence']}")

                                st.markdown(f"**Source:** {m['source_sentence']}")

                                st.caption(f"File: {m['source_file']} | Page: {m['page']} | Score: {m['similarity_score']:.3f} | Chars: {m['char_range']}")

                                st.markdown("---")





else:

    st.title("Tender Analyzer — Context Understanding Mode")

    with st.sidebar:

        st.header("SUM Settings")

        gem_key = st.text_input("Google API Key", type="password", key="gka")

        s_obj = st.radio("Objective", ["General Summary", "Compliance Matrix", "Risk Assessment", "Entity Dashboard", "Ambiguity Scrutiny"], key="sob")

        s_model = st.text_input("Model", value="gemini-2.5-flash", key="smk")

        s_rpm = st.number_input("Target RPM", min_value=1, max_value=4, value=4, key="srp")

        s_batch = st.number_input("Batch size", min_value=10, max_value=50, value=10, key="sbt")

        s_max_tk = st.number_input("Max tokens/chunk", min_value=500, max_value=8000, value=8000, key="smt")

        s_over = st.number_input("Overlap sents", value=5, key="sov")

        s_prompt = st.text_area("Extra prompts", key="sep")



    f_sum = st.file_uploader("Upload PDFs", type="pdf", accept_multiple_files=True, key="sf")

    q_sum = st.text_area("Queries (one per line):", key="sq")



    if st.button("Analyze", key="sab"):

        if not f_sum:

            st.error("Please upload at least one PDF file")

            st.stop()

        if not gem_key:

            st.error("Please provide a Google API Key in the sidebar")

            st.stop()

        if not q_sum or not q_sum.strip():

            st.error("Please provide at least one query")

            st.stop()



        # Validate file sizes

        for f in f_sum:

            file_size_mb = len(f.getvalue()) / (1024 * 1024)

            if file_size_mb > MAX_FILE_SIZE_MB:

                st.error(f"File {f.name} exceeds maximum size of {MAX_FILE_SIZE_MB}MB (size: {file_size_mb:.1f}MB)")

                st.stop()



        client = genai.Client(api_key=gem_key)

        all_p, all_f = [], []

        temp_files = []



        try:

            for f in f_sum:

                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as t:

                    temp_files.append(t.name)

                    t.write(f.getbuffer())

                    t.flush()

                    p, fl = extract_pages_sum(t.name)

                    all_p.extend(p)

                    all_f.extend(fl)



            if not all_p:

                st.error("No content could be extracted from the uploaded files")

                st.stop()



            chunks = semantic_chunk_pages_sum(all_p, all_f, max_tok=s_max_tk, overlap=s_over)



            if not chunks:

                st.error("No chunks could be created from the extracted content")

                st.stop()

        finally:

            # Clean up temporary files

            for temp_file in temp_files:

                try:

                    os.unlink(temp_file)

                except Exception:

                    pass



        async def run_sum():

            ms, mi, rs, ri = get_sum_prompts(s_obj)

            results = []

            instr = s_prompt.strip() + "\n\n" if s_prompt.strip() else ""

            batches = [chunks[i:i+s_batch] for i in range(0, len(chunks), s_batch)]

            for q in q_sum.splitlines():

                if not q.strip(): 

                    continue

                if len(q) > MAX_QUERY_LENGTH:

                    st.warning(f"Skipping query (too long): {q[:100]}...")

                    continue

                t0 = time.time()

                m_tasks = []

                for b in batches:

                    b_txt = '\n'.join([c['text'] for c in b])

                    m_tasks.append(call_gemini_json_sum_async(client, ms, f"{instr}{mi}\nQ:{q}\nC:{b_txt}", s_model, s_rpm))

                mapped = await asyncio.gather(*m_tasks)

                t_m = time.time()

                red = await call_gemini_json_sum_async(client, rs, f"{instr}{ri}\nQ:{q}\nD:{json.dumps(mapped)}", s_model, s_rpm)

                results.append({"query":q, "result":red, "map_t":t_m-t0, "red_t":time.time()-t_m, "mapped":mapped})

            return results



        final_res = asyncio.run(run_sum())

        for r in final_res:

            st.subheader(f"Query: {r['query']}")

            st.caption(f"⏱️ Map: {r['map_t']:.2f}s | Reduce: {r['red_t']:.2f}s")



            # Check for errors in result

            if isinstance(r["result"], dict) and "error" in r["result"]:

                st.error(f"Error processing query: {r['result']['error']}")

                continue



            st.json(r["result"])

            conf = compute_confidence_score_sum(r['mapped'], r['result'], r['query'])

            st.info(f"Confidence: {conf['overall_confidence']:.2%}")


            
            # Convert results to table format and display
            df = convert_result_to_dataframe(r["result"], s_obj)
            
            if df is not None and not df.empty:
                st.markdown("---")
                st.markdown("### 📊 Results Table")
                st.dataframe(df, use_container_width=True)
                
                # Export buttons
                col1, col2 = st.columns(2)
                
                with col1:
                    excel_data = export_to_excel(df, r['query'], s_obj)
                    if excel_data:
                        safe_objective = sanitize_filename(s_obj)
                        safe_query = sanitize_filename(r['query'])
                        st.download_button(
                            label="📥 Download Excel",
                            data=excel_data,
                            file_name=f"{safe_objective}_{safe_query}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"excel_{r['query']}"
                        )
                
                with col2:
                    word_data = export_to_word(df, r['query'], s_obj)
                    if word_data:
                        safe_objective = sanitize_filename(s_obj)
                        safe_query = sanitize_filename(r['query'])
                        st.download_button(
                            label="📥 Download Word",
                            data=word_data,
                            file_name=f"{safe_objective}_{safe_query}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"word_{r['query']}"
                        )
            elif df is not None and df.empty and s_obj == "Entity Dashboard":
                # Special handling for empty Entity Dashboard
                st.markdown("---")
                st.markdown("### 📊 Entity Dashboard")
                st.info(
                    "ℹ️ No entities were found in the analyzed document sections. This could mean:\n"
                    "- The document sections don't contain recognizable entities\n"
                    "- The query may need to be adjusted\n"
                    "- Try analyzing different sections or the entire document"
                )
                
                # Show entity count summary (will display zeros for all categories)
                if isinstance(r["result"], dict) and "entity_count" in r["result"]:
                    st.markdown("**Entity Count Summary:**")
                    entity_count_df = pd.DataFrame([r["result"]["entity_count"]])
                    st.dataframe(entity_count_df, use_container_width=True)


            # Find pages for preview

            pgs = []

            def _f(o):

                if isinstance(o, dict):

                    for k, v in o.items():

                        if k in ["page", "pages", "citations"]:

                            if isinstance(v, list):

                                for x in v:

                                    if isinstance(x, int):

                                        pgs.append(x)

                            elif isinstance(v, int):

                                pgs.append(v)

                        else:

                            _f(v)

                elif isinstance(o, list):

                    for x in o:

                        _f(x)

            _f(r['result'])



            # Display page-level citation tracking with text excerpts
            if pgs:
                render_page_level_citations_sum(chunks, sorted(set(pgs)))


            if pgs and f_sum:

                try:

                    pdf_doc = fitz.open(stream=f_sum[0].getvalue(), filetype="pdf")

                    render_citation_preview_sum(pdf_doc, [{"page": p} for p in set(pgs)])

                    pdf_doc.close()

                except Exception as e:

                    st.error(f"Failed to load PDF for preview: {e}")
